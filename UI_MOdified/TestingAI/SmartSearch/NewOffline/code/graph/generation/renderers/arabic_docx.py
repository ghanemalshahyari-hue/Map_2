"""graph/generation/renderers/arabic_docx.py

Arabic ``.docx`` renderer for Phase 3. Ported from the user's prior
generator (``/Users/hextechkraken/Desktop/ToTransfer/New Text
Document.txt``). Behaviour-preserving: constants, margins, kashida
logic, bidi handling, numbering cycles, and table styling match the
old code's rendered output. Do NOT simplify. Prior simplification
attempts broke formatting — see
``referencedocs/19_phase3_arabic_renderer.md`` §1.

Scope change vs the old code (the ONE allowed non-cosmetic change,
per ref 19 §4): the module-level ``LEVEL_COUNTERS`` dict is moved
into a per-document :class:`ArabicDocumentContext` so four documents
can render in any order, in any combination, without cross-document
counter leakage. ``reset_lower_counters`` now operates on the
context's counters; every ``add_level_*`` call takes the context as
first argument.

Other cosmetic tidying (ref 19 §3.3):

- Dropped top-level debug ``print("level1")`` / ``print("level2")``
  / ... inside ``SPLITTER``.
- Hostile ``except ImportError: exit(1)`` guard at module import
  replaced with a deferred ``ModuleNotFoundError`` raised at first
  call site.
- Single de-duplicated ``from docx.enum.text import WD_ALIGN_PARAGRAPH``
  (the old file imports it twice).
- ``except Exception: pass`` blocks narrowed to specific OOXML
  failure modes where possible.
- Type hints added on the port targets; originals remain inconsistent.

Module entry point for the rest of the generator:
:func:`render_document` walks a
:class:`~graph.generation.assembler.GeneratedDocument` alongside
its :class:`~graph.generation.template_loader.Template` and writes
one ``.docx`` file per call.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_PARAGRAPH_ALIGNMENT,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph as DocxParagraph

try:
    import PIL.ImageFont as _ImageFont  # type: ignore[import-not-found]
    import arabic_reshaper as _arabic_reshaper  # type: ignore[import-not-found]
    _HAS_ARABIC_TOOLCHAIN = True
except ImportError:  # pragma: no cover — renderer still works without kashida
    _ImageFont = None  # type: ignore[assignment]
    _arabic_reshaper = None  # type: ignore[assignment]
    _HAS_ARABIC_TOOLCHAIN = False


# =============================================================================
# 1. CONFIGURATION (verbatim from old code)
# =============================================================================

MARGIN_CM = Cm(2.5)
LINE_SPACING_CM = Cm(0.75)
TAB_SIZE_CM = Cm(1.5)
FONT_NAME = "Arial"
FONT_SIZE_PT = Pt(14)
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)

# Full 28-letter Arabic cycle used by add_level_two / add_level_four.
ARABIC_LETTERS = [
    "أ", "ب", "جـ", "د", "هـ", "و", "ز", "حـ", "ط", "ي", "ك", "ل",
    "م", "ن", "س", "ع", "ف", "ص", "ق", "ر", "ش", "ت", "ث", "خ",
    "ذ", "ض", "ظ", "غ",
]
# First twelve — used for appendix sub-lists (ملاحق).
ARABIC_LETTERS_ML = ["أ", "ب", "جـ", "د", "هـ", "و", "ز", "حـ", "ط", "ي", "ك", "ل"]
# Letters 13–28 — used for overlay sub-lists (شفافيات).
ARABIC_LETTERS_SHFAF = [
    "م", "ن", "س", "ع", "ف", "ص", "ق", "ر", "ش", "ت", "ث", "خ",
    "ذ", "ض", "ظ", "غ",
]


# =============================================================================
# 2. CONTEXT (new — replaces module-level LEVEL_COUNTERS per ref 19 §4)
# =============================================================================

@dataclass
class ArabicDocumentContext:
    """Per-document state for the renderer.

    Replaces the old code's module-level ``LEVEL_COUNTERS`` dict so
    four documents can render in any order without cross-document
    counter leakage. Counter keys match the old names for minimal
    drift (``level1`` … ``level10``). The behaviour of any single
    document's rendering is identical to the old code because
    ``reset_lower_counters`` was already called at the top of each
    document in the old ``generate_document``; moving the state off
    module-level only affects the leak, not rendered output.
    """

    document: Document
    counters: dict[str, int] = field(default_factory=lambda: {
        f"level{n}": 0 for n in range(1, 11)
    })

    @classmethod
    def new(cls) -> "ArabicDocumentContext":
        """Create a fresh context with a new Document + zero counters."""
        return cls(document=Document())


def reset_lower_counters(ctx: ArabicDocumentContext, start_level: int) -> None:
    """Zero out counters at ``start_level`` and below.

    Verbatim behaviour of the old module-level function, now scoped
    to ``ctx.counters``.
    """
    for level in range(start_level, 9):
        ctx.counters[f"level{level}"] = 0


# =============================================================================
# 3. SMALL UTILS (verbatim)
# =============================================================================

def get_arabic_letter(index: int) -> str:
    if index <= 0:
        raise ValueError("Index must be positive.")
    return ARABIC_LETTERS[(index - 1) % len(ARABIC_LETTERS)]


def get_arabic_letter_ML(index: int) -> str:
    if index <= 0:
        raise ValueError("Index must be positive.")
    return ARABIC_LETTERS_ML[(index - 1) % len(ARABIC_LETTERS_ML)]


def get_arabic_letter_SHFAF(index: int) -> str:
    if index <= 0:
        raise ValueError("Index must be positive.")
    return ARABIC_LETTERS_SHFAF[(index - 1) % len(ARABIC_LETTERS_SHFAF)]


def normalize_text(text: str) -> str:
    return re.sub(
        r"\s+",
        " ",
        text.replace("\n", " ").replace("\r", " ").replace("ـ", ""),
    ).strip()


def add_full_stop(sentence: str) -> str:
    if not isinstance(sentence, str) or not sentence.strip():
        return sentence
    sentence = sentence.strip()
    if sentence[-1] not in {".", "!", "?", "؟", "؛", "،", ":", ";", ")", "]"}:
        sentence += "."
    return sentence


# =============================================================================
# 4. RUN / PARAGRAPH FORMATTING (verbatim from old code §4 — critical,
# do NOT simplify — fix_cs_formatting_run and force_rtl_paragraph are
# belt-and-suspenders against python-docx dropping bidi markers)
# =============================================================================

def fix_cs_formatting_run(
    run,
    font_name: str = FONT_NAME,
    font_size_pt=FONT_SIZE_PT,
    bold: bool = False,
    underline: bool = False,
    tabs: int = 0,
    rtl: bool = True,
) -> None:
    font = run.font
    font.name = font_name
    font.size = font_size_pt
    font.cs_size = font_size_pt
    font.bold = bold
    font.cs_bold = bold
    font.underline = underline
    font.rtl = True
    font.complex_script = True
    rPr = run.element.get_or_add_rPr()
    sz = OxmlElement("w:sz")
    szCs = OxmlElement("w:szCs")
    sz.set(qn("w:val"), str(int(font_size_pt.pt * 2)))
    szCs.set(qn("w:val"), str(int(font_size_pt.pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:bidi"), "ar-SA")
    rPr.append(lang)


def force_rtl_paragraph(paragraph, tabs: int = 0) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)


def add_paragraph(
    document: Document,
    text: str,
    bold: bool = False,
    underline: bool = False,
    rtl: bool = True,
    tabs: int = 0,
    middle_alignment: bool = False,
    space_before: float = 0.0,
    space_after: float = 0.0,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING_CM
    fmt.space_before = Cm(space_before)
    fmt.space_after = Cm(space_after)
    fmt.right_to_left = rtl
    indent = TAB_SIZE_CM * tabs
    if rtl:
        fmt.right_indent = indent
        fmt.left_indent = Cm(0)
    else:
        fmt.left_indent = indent
        fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
        if middle_alignment
        else (WD_PARAGRAPH_ALIGNMENT.RIGHT if rtl else WD_PARAGRAPH_ALIGNMENT.LEFT)
    )
    run = paragraph.add_run(text)
    fix_cs_formatting_run(
        run=run,
        font_name=FONT_NAME,
        font_size_pt=FONT_SIZE_PT,
        bold=bold,
        underline=underline,
    )
    return paragraph


def correct_indentation(
    document: Document,
    body: str,
    prefix_text: str,
    tabs_before: int = 0,
    underline: bool = False,
    space_before: float = 0.0,
    space_after: float = 0.0,
    tabs_after: int = 1,
    tabs: int = 0,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    format_ = paragraph.paragraph_format
    format_.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    format_.line_spacing = LINE_SPACING_CM
    format_.space_before = Cm(space_before)
    format_.space_after = Cm(space_after)
    format_.right_to_left = True
    format_.rtl = True
    format_.right_indent = Cm(0)
    format_.left_indent = Cm(0)
    format_.first_line_indent = Cm(0)
    format_.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    force_rtl_paragraph(paragraph)
    indent = TAB_SIZE_CM * tabs_before
    format_.left_indent = indent
    tab_stops = format_.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(
        indent + TAB_SIZE_CM,
        alignment=WD_TAB_ALIGNMENT.LEFT,
        leader=WD_TAB_LEADER.SPACES,
    )
    prefix_run = paragraph.add_run(prefix_text + "\t")
    prefix_run.underline = False
    fix_cs_formatting_run(prefix_run, FONT_NAME, FONT_SIZE_PT, bold=False)
    body_run = paragraph.add_run(body)
    fix_cs_formatting_run(
        body_run, FONT_NAME, FONT_SIZE_PT, bold=False, underline=underline, tabs=tabs
    )


# =============================================================================
# 5. NUMBERED LEVELS (verbatim — counters now live on ctx, per §4 of ref 19)
# =============================================================================

def add_level_one(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
    space_after: float = 0.0,
) -> None:
    ctx.counters["level1"] += 1
    reset_lower_counters(ctx, 2)
    prefix = f"{ctx.counters['level1']}."
    correct_indentation(
        ctx.document,
        body,
        prefix,
        tabs_before=0,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        space_after=space_after,
    )


def add_level_two(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
    space_after: float = 0.0,
) -> None:
    ctx.counters["level2"] += 1
    reset_lower_counters(ctx, 3)
    letter = get_arabic_letter(ctx.counters["level2"])
    correct_indentation(
        ctx.document,
        body,
        f"{letter}.",
        tabs_before=1,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=1,
        space_after=space_after,
    )


def add_level_three(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
    space_after: float = 0.0,
) -> None:
    ctx.counters["level3"] += 1
    reset_lower_counters(ctx, 4)
    prefix = f"({ctx.counters['level3']})"
    correct_indentation(
        ctx.document,
        body,
        prefix,
        tabs_before=2,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=2,
        space_after=space_after,
    )


def add_level_four(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
) -> None:
    ctx.counters["level4"] += 1
    reset_lower_counters(ctx, 5)
    letter = get_arabic_letter(ctx.counters["level4"])
    correct_indentation(
        ctx.document,
        body,
        f"({letter})",
        tabs_before=3,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=3,
    )


def add_level_five(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
) -> None:
    ctx.counters["level5"] += 1
    reset_lower_counters(ctx, 6)
    letter = get_arabic_letter(ctx.counters["level5"])
    correct_indentation(
        ctx.document,
        body,
        f"({letter} {letter})",
        tabs_before=4,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=4,
    )


def add_level_one_ML(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
) -> None:
    ctx.counters["level6"] += 1
    reset_lower_counters(ctx, 3)
    if ctx.counters["level6"] <= 0:
        ctx.counters["level6"] = 10
    letter = get_arabic_letter_ML(ctx.counters["level6"])
    correct_indentation(
        ctx.document,
        body,
        f"{letter}.",
        tabs_before=0,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=1,
    )


def add_level_one_SHFAF(
    ctx: ArabicDocumentContext,
    body: str,
    underline: bool = False,
    space_before: float = 0.0,
) -> None:
    ctx.counters["level7"] += 1
    reset_lower_counters(ctx, 3)
    if ctx.counters["level7"] <= 0:
        ctx.counters["level7"] = 10
    letter = get_arabic_letter_SHFAF(ctx.counters["level7"])
    correct_indentation(
        ctx.document,
        body,
        f"{letter}.",
        tabs_before=0,
        tabs_after=1,
        underline=underline,
        space_before=space_before,
        tabs=1,
    )


def append_to_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    underline: bool = False,
    rtl: bool = True,
) -> None:
    para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    run = para.add_run(text)
    fix_cs_formatting_run(run, FONT_NAME, FONT_SIZE_PT, bold=False, underline=underline)
    if underline:
        run.underline = True
    para.paragraph_format.right_to_left = rtl


def SPLITTER(text: str, ctx: ArabicDocumentContext) -> list[str]:
    """Split a multi-line text that already carries embedded
    Arabic numbering into nested level-N paragraphs.

    Matches old behaviour but:

    - Takes ctx instead of document (counters are per-document).
    - Drops the debug print()s from the old code (ref 19 §3.2).
    """
    sentences = [s for s in text.splitlines() if s.strip()]
    for s in sentences:
        if s.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")):
            add_level_one(ctx, s.split(".", 1)[1].strip(), underline=False)
        elif s.startswith(("أ.", "ب.", "جـ.", "د.", "هـ.", "و.", "ز.")):
            add_level_two(ctx, s.split(".", 1)[1].strip(), underline=False)
        elif s.startswith(("(1)", "(2)", "(3)", "(4)", "(5)", "(6)", "(7)", "(8)")):
            add_level_three(ctx, s.split(")", 1)[1].strip(), underline=False)
        elif s.startswith((
            "(أ)", "(ب)", "(جـ)", "(د)", "(هـ)", "(و)", "(ز)",
            "( أ )", "( ب )", "( جـ )", "( د )", "( هـ )", "( و )", "( ز )",
            "( أ)", "( ب)", "( جـ)", "( د)", "( هـ)", "( و)", "( ز)",
            "(أ )", "(ب )", "(جـ )", "(د )", "(هـ )", "(و )", "(ز )",
            "\t(أ)", "\t(ب)", "\t(جـ)", "\t(د)", "\t(هـ)", "\t(و)", "\t(ز)",
            "\t( أ )", "\t( ب )", "\t( جـ )", "\t( د )", "\t( هـ )", "\t( و )", "\t( ز )",
            "\t(أ )", "\t(ب )", "\t(جـ )", "\t(د )", "\t(هـ )", "\t(و )", "\t(ز )",
            "\t( أ)", "\t( ب)", "\t( جـ)", "\t( د)", "\t( هـ)", "\t( و)", "\t( ز)",
        )):
            add_level_four(ctx, s.split(")", 1)[1].strip())
        elif s.startswith((
            "(أ أ)", "(ب ب)", "(جـ جـ)", "(د د)", "(هـ هـ)", "(و و)", "(ز ز)",
            "( أ أ )", "( ب ب )", "( جـ جـ )", "( د د )", "( هـ هـ )", "( و و )", "( ز ز )",
            "(أأ)", "(بب)", "(جـجـ)", "(دد)", "(هـهـ)", "(وو)", "(زز)",
            "( أ أ)", "( ب ب)", "( جـ جـ)", "( د د)", "( هـ هـ)", "( و و)", "( ز ز)",
            "(أ أ )", "(ب ب )", "(جـ جـ )", "(د د )", "(هـ هـ )", "(و و )", "(ز ز )",
            "\t(أ أ)", "\t(ب ب)", "\t(جـ جـ)", "\t(د د)", "\t(هـ هـ)", "\t(و و)", "\t(ز ز)",
            "\t( أ أ )", "\t( ب ب )", "\t( جـ جـ )", "\t( د د )", "\t( هـ هـ )", "\t( و و )", "\t( ز ز )",
            "\t(أأ)", "\t(بب)", "\t(جـجـ)", "\t(دد)", "\t(هـهـ)", "\t(وو)", "\t(زز)",
            "\t( أ أ)", "\t( ب ب)", "\t( جـ جـ)", "\t( د د)", "\t( هـ هـ)", "\t( و و)", "\t( ز ز)",
            "\t(أ أ )", "\t(ب ب )", "\t(جـ جـ )", "\t(د د )", "\t(هـ هـ )", "\t(و و )", "\t(ز ز )",
        )):
            add_level_five(ctx, s.split(")", 1)[1].strip())
    return sentences


# =============================================================================
# 6. TABLE (verbatim)
# =============================================================================

def add_table(
    document: Document,
    rows: list[tuple[str, ...]],
    headers: list[str],
    rows_length: int = -1,
    column_length: int = -1,
    enable_autofit: bool = True,
    custom_widths: list | None = None,
) -> None:
    if rows_length < 0:
        num_rows = len(rows) + 1  # +1 for header
        table = document.add_table(rows=num_rows, cols=5)
    else:
        num_rows = len(rows) + 1
        table = document.add_table(rows=num_rows, cols=rows_length)
    table.style = None
    if enable_autofit:
        tbl_pr = table._element.tblPr
        if tbl_pr is None:
            tbl_pr = table.element.get_or_add_tblPr()
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "autofit")
    else:
        table.autofit = False
        if custom_widths is None:
            custom_widths = [Cm(1.0), Cm(3.75), Cm(3.75), Cm(3.75), Cm(3.75), Cm(4.0)]
        for i, width in enumerate(custom_widths[: len(table.columns)]):
            table.columns[i].width = width
    try:
        tbl_pr = table._element.get_or_add_tblPr()
        tc_mar = OxmlElement("w:tblInd")
        tc_mar.set(qn("w:w"), str(int(1.5 * 1440)))
        tc_mar.set(qn("w:type"), "dxa")
        tbl_pr.append(tc_mar)
    except (AttributeError, KeyError, ValueError):
        # OOXML surface here is finicky — swallow known failure modes
        # but let real bugs propagate.
        pass
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
        tbl_pr = table._element.tblPr
        if tbl_pr is None:
            tbl_pr = table._element.get_or_add_tblPr()
        bidi_visual = tbl_pr.get_or_add_bidiVisual()
        bidi_visual.set(qn("w:val"), "1")
    except (AttributeError, KeyError, ValueError):
        pass

    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        paragraph = cell.paragraphs[0]
        if hasattr(paragraph, "clear"):
            paragraph.clear()
        else:
            for r in paragraph.runs:
                r.clear()
        paragraph.style = document.styles["Normal"]
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = LINE_SPACING_CM
        fmt.space_before = Cm(1.0)
        fmt.space_after = Cm(0)
        fmt.right_to_left = True
        fmt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fmt.right_indent = Cm(0)
        fmt.left_indent = Cm(0)
        fmt.first_line_indent = Cm(0)
        try:
            tc_pr = cell._element.get_or_add_tcPr()
            existing_shd = tc_pr.find(qn("w:shd"))
            if existing_shd is not None:
                tc_pr.remove(existing_shd)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "DDEBF7")
            tc_pr.append(shading)
        except (AttributeError, KeyError, ValueError):
            pass
        run = paragraph.add_run(header_text)
        force_rtl_paragraph(paragraph)
        fix_cs_formatting_run(
            run, FONT_NAME, FONT_SIZE_PT, bold=False, underline=False, rtl=True
        )
    for row_idx, a in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        cell_data = a
        for i, cell_text in enumerate(cell_data):
            cell = row_cells[i]
            paragraph = cell.paragraphs[0]
            if hasattr(paragraph, "clear"):
                paragraph.clear()
            else:
                for r in paragraph.runs:
                    r.clear()
            paragraph.style = document.styles["Normal"]
            fmt = paragraph.paragraph_format
            fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            fmt.line_spacing = LINE_SPACING_CM
            fmt.space_before = Cm(0)
            fmt.space_after = Cm(0)
            fmt.right_to_left = True
            fmt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            fmt.right_indent = Cm(0)
            fmt.left_indent = Cm(0)
            fmt.first_line_indent = Cm(0)
            try:
                shading = cell._element.get_or_add_tcPr().get_or_add_shd()
                shading.set(qn("w:fill"), "FFFFFF")
                shading.set(qn("w:val"), "clear")
            except (AttributeError, KeyError, ValueError):
                pass
            display_text = cell_text if cell_text or i == 0 else "......"
            run = paragraph.add_run(display_text)
            force_rtl_paragraph(paragraph)
            fix_cs_formatting_run(run, FONT_NAME, FONT_SIZE_PT, bold=False, rtl=True)


# =============================================================================
# 7. HEADER + PAGE-NUMBER (verbatim — kashida stretcher preserved byte-for-byte)
# =============================================================================

def add_arabic_header(
    doc: Document,
    copy_number_placeholder: str = "النسخـــــــــــــــــــــــــة رقـــــــــــم..... مــــــن.....",
    organization: str = "وزارة الصحة",
    main_unit: str = "وحدة التمريض",
    sub_unit: str = "مجموعة الصحة",
    place: str = "",
    date_hijri: str = "",
    date_greg: str = "",
    letter_ref_number: str = "",
    target_length_chars: int = 65,
    max_kashidas_per_word: int = 30,
    cell_width_right: Any = Cm(15.0),
    cell_width_left: Any = Cm(1.0),
) -> None:
    _ARIAL_TTF = "arial.ttf"
    HAS_KASHIDA = _HAS_ARABIC_TOOLCHAIN

    if HAS_KASHIDA:
        def rendered_width_px(text: str, font_pt: int = 14, dpi: int = 96) -> float:
            px_size = max(1, int(round(font_pt * dpi / 72)))
            try:
                font = _ImageFont.truetype(_ARIAL_TTF, px_size)
            except Exception:
                font = _ImageFont.load_default()
            if hasattr(font, "getlength"):
                return float(font.getlength(text))
            bbox = font.getbbox(text)
            return float(bbox[2] - bbox[0])

        def _is_arabic_letter(ch: str) -> bool:
            return 0x0600 <= ord(ch) <= 0x06FF

        def calculate_kashida(text: str) -> str:
            if not text.strip() or len(text) < 2:
                return text
            target_px = (8.0 / 2.54) * 99 * 0.99
            if rendered_width_px(text) >= target_px:
                return text
            valid_indices = []
            for i in range(len(text) - 1):
                a, b = text[i], text[i + 1]
                if " " in (a, b):
                    continue
                if not (_is_arabic_letter(a) and _is_arabic_letter(b)):
                    continue
                if a in "اأإآدذرزو.ىؤء":
                    continue
                if a == "ل" and b in "أإآا":
                    continue
                if b == "ء":
                    continue
                valid_indices.append(i)
            if not valid_indices:
                return text
            counts = {i: 0 for i in valid_indices}
            ptr = 0
            for _ in range(500):
                idx = valid_indices[ptr]
                counts[idx] += 1
                result_chars: list[str] = []
                for i, ch in enumerate(text):
                    result_chars.append(ch)
                    if i in counts and counts[i] > 0:
                        result_chars.append("ـ" * counts[i])
                candidate = "".join(result_chars)
                if rendered_width_px(candidate) > target_px:
                    counts[idx] -= 1
                    break
                ptr = (ptr + 1) % len(valid_indices)
            result_chars = []
            for i, ch in enumerate(text):
                result_chars.append(ch)
                if i in counts and counts[i] > 0:
                    result_chars.append("ـ" * counts[i])
            return "".join(result_chars)

        def stretch_first_line(text: str, target_chars: int = 65, max_kashidas: int = 30) -> str:
            words = text.split()
            if not words:
                return text
            stretched_words: list[str] = []
            for word in words:
                if len(word) < 2:
                    stretched_words.append(word)
                    continue
                mid = len(word) // 2
                if word[mid - 1] not in "اأإآد.ذرزوىؤء" and word[mid] != "ء":
                    kashida_count = min(
                        max_kashidas, max(1, (target_chars - len(text)) // len(words) + 2)
                    )
                    word = word[:mid] + "ـ" * kashida_count + word[mid:]
                stretched_words.append(word)
            return " ".join(stretched_words)
    else:
        def calculate_kashida(text: str) -> str:
            return text

        def stretch_first_line(text: str, target_chars: int = 65, max_kashidas: int = 30) -> str:
            return text

    def set_cell_width(cell, width: Any) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        child = tcPr.first_child_found_in("w:tcW")
        tcW = child if child is not None else OxmlElement("w:tcW")
        if hasattr(width, "cm"):
            w_twips = int(width.cm * 567)
        elif hasattr(width, "pt"):
            cm_value = width.pt * 2.54 / 72
            w_twips = int(cm_value * 567)
        else:
            w_twips = int(float(width.cm) * 567)
        tcW.set(qn("w:w"), str(w_twips))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    def set_table_column_widths(table, width_left: Any, width_right: Any) -> None:
        for row in table.rows:
            set_cell_width(row.cells[0], width_left)
            set_cell_width(row.cells[1], width_right)

    def tighten_cell(cell) -> None:
        for p in cell.paragraphs:
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def process_cell(cell, text: str, underline: bool = False, use_kashida: bool = True) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.text = ""
        if use_kashida and text == copy_number_placeholder:
            text = stretch_first_line(text, target_chars=target_length_chars, max_kashidas=max_kashidas_per_word)
        elif use_kashida:
            text = calculate_kashida(text)
        try:
            reshaped = _arabic_reshaper.reshape(text) if HAS_KASHIDA else text
        except Exception:
            reshaped = text
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = LINE_SPACING_CM
        fmt.right_to_left = True
        fmt.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        run = paragraph.add_run(reshaped)
        fix_cs_formatting_run(run, bold=False, underline=underline)

    table = doc.add_table(rows=8, cols=2)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.75)
    set_table_column_widths(table, cell_width_left, cell_width_right)
    items = [
        copy_number_placeholder,
        organization,
        main_unit,
        sub_unit,
        place,
        date_hijri,
        date_greg,
        letter_ref_number,
    ]
    for i, txt in enumerate(items):
        process_cell(
            table.rows[i].cells[1],
            txt,
            underline=(txt == sub_unit),
            use_kashida=(txt != letter_ref_number),
        )
        tighten_cell(table.rows[i].cells[1])

    last_paragraph = doc.add_paragraph()
    last_paragraph.paragraph_format.space_after = Cm(0.3)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = paragraph.add_run("-")
    r1.font.name = FONT_NAME
    r1.font.size = FONT_SIZE_PT
    r1.font.bold = True
    r1.font.underline = True
    r1.font.rtl = True
    r1.font.complex_script = True
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    for element in (fld_char_begin, instr_text, fld_char_end):
        run._r.append(element)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_PT
    run.font.bold = True
    run.font.underline = True
    run.font.rtl = True
    run.font.complex_script = True
    r3 = paragraph.add_run("-")
    r3.font.name = FONT_NAME
    r3.font.size = FONT_SIZE_PT
    r3.font.bold = True
    r3.font.underline = True
    r3.font.rtl = True
    r3.font.complex_script = True


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_PT
    font.cs_size = FONT_SIZE_PT
    font.rtl = True
    font.complex_script = True
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING_CM
    pf.space_before = Cm(0)
    pf.space_after = Cm(0)
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pf.right_to_left = True
    section = document.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = MARGIN_CM
    section.bottom_margin = MARGIN_CM
    section.left_margin = MARGIN_CM
    section.right_margin = MARGIN_CM
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    for p in first_header.paragraphs:
        p.text = ""
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    for p in first_footer.paragraphs:
        p.text = ""
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ""
    header_para = header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_before = Cm(1)
    header_para.paragraph_format.space_after = Cm(1)
    header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    header_para.paragraph_format.line_spacing = Pt(18)
    run = header_para.add_run("سري للغاية")
    fix_cs_formatting_run(run, font_name=FONT_NAME, font_size_pt=FONT_SIZE_PT, bold=True, underline=False)
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""
    num_para = footer.add_paragraph()
    num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(num_para)
    siri_para = footer.add_paragraph("سري للغاية")
    siri_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = siri_para.paragraph_format
    fmt.space_before = Cm(0)
    fmt.space_after = Cm(2.5)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    for run in siri_para.runs:
        fix_cs_formatting_run(run, font_name=FONT_NAME, font_size_pt=FONT_SIZE_PT, bold=True, underline=False)


def configure_last_page_section(document: Document) -> None:
    first_section = document.sections[0]
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    for attr in (
        "page_width", "page_height",
        "top_margin", "bottom_margin", "left_margin", "right_margin",
        "header_distance", "footer_distance",
    ):
        setattr(new_section, attr, getattr(first_section, attr))
    new_section.different_first_page_header_footer = False
    new_section.header.is_linked_to_previous = False
    footer = new_section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    def _append_footer_para(text: str = "", bold: bool = True, underline: bool = False, space_after_cm: float = 0.0):
        p_el = OxmlElement("w:p")
        footer._element.append(p_el)
        para = DocxParagraph(p_el, footer)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = para.paragraph_format
        fmt.space_before = Cm(0)
        fmt.space_after = Cm(space_after_cm)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(18)
        if text:
            run = para.add_run(text)
            fix_cs_formatting_run(run, font_name=FONT_NAME, font_size_pt=FONT_SIZE_PT, bold=bold, underline=underline)
        return para

    lp = footer.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Cm(0)
    lp.paragraph_format.space_after = Cm(0)
    lp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    lp.paragraph_format.line_spacing = Pt(18)
    run = lp.add_run("الصفحة الأخيرة")
    fix_cs_formatting_run(run, font_name=FONT_NAME, font_size_pt=FONT_SIZE_PT, bold=True)
    num_para = _append_footer_para()
    add_page_number(num_para)
    _append_footer_para("سري للغاية", bold=True, underline=True, space_after_cm=2.5)


# =============================================================================
# 8. TEMPLATE-DRIVEN ORCHESTRATION (new — wires renderer to the rest of M2)
# =============================================================================

def _extra_get(entry, key: str) -> Any:
    """Safe read of a pydantic-extra attribute on a StructureEntry."""
    extras = entry.model_extra or {}
    return extras.get(key)


def _render_title(ctx: ArabicDocumentContext, entry) -> None:
    text = _extra_get(entry, "text") or ""
    underline = bool(_extra_get(entry, "underline"))
    alignment = _extra_get(entry, "alignment") or "center"
    middle = alignment == "center"
    add_paragraph(
        ctx.document,
        text,
        bold=True,
        underline=underline,
        rtl=True,
        middle_alignment=middle,
        space_before=0.5,
        space_after=0.5,
    )


def _render_heading(ctx: ArabicDocumentContext, heading: dict | None) -> None:
    if not heading:
        return
    text = heading.get("text", "") if isinstance(heading, dict) else getattr(heading, "text", "")
    if not text:
        return
    underline = bool(heading.get("underline") if isinstance(heading, dict) else getattr(heading, "underline", False))
    space_before = float(heading.get("space_before_cm", 0.5) if isinstance(heading, dict) else getattr(heading, "space_before_cm", 0.5))
    space_after = float(heading.get("space_after_cm", 0.0) if isinstance(heading, dict) else getattr(heading, "space_after_cm", 0.0))
    add_paragraph(
        ctx.document,
        text,
        bold=True,
        underline=underline,
        rtl=True,
        middle_alignment=False,
        space_before=space_before,
        space_after=space_after,
    )


def _layout_numbered_fields(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Emit each field of the class as a level-1 numbered paragraph.

    Uses the YAML ``label_ar`` for the field label when declared
    (otherwise falls back to the field name). If the value carries
    embedded Arabic numbering (e.g. "1. …\\n2. …"), the SPLITTER
    pass decomposes it into nested levels.
    """
    template = generated.template
    pydantic_instance = generated.sections[schema_name]
    fields = template.schemas[schema_name].fields
    for fname, spec in fields.items():
        label = (getattr(spec, "label_ar", None) or fname)
        value = getattr(pydantic_instance, fname, "")
        body = f"{label}: {value}" if value else label
        add_level_one(ctx, body)


def _reset_level_one_counter(ctx: ArabicDocumentContext) -> None:
    """Reset the level-1 numbering so a new titled sub-section
    restarts at ``1.``. Also clears levels 2+ via the usual
    :func:`reset_lower_counters` helper.

    Used by the ``staff_sections`` layout (Doc 2) so each staff
    estimate — الاستخبارات / العمليات / الموارد البشرية / الإمداد —
    numbers its fields independently.
    """
    ctx.counters["level1"] = 0
    reset_lower_counters(ctx, 2)


def _layout_staff_sections(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Doc 2's layout: each Pydantic class becomes a self-contained
    numbered sub-section.

    Behaviour: reset level-1 to 0 at the top of the section so the
    first field is rendered as ``1.`` regardless of what came above,
    then fall through to :func:`_layout_numbered_fields` for the
    actual field emission. The heading line itself is rendered by
    :func:`_render_heading` before this is called.
    """
    _reset_level_one_counter(ctx)
    _layout_numbered_fields(ctx, schema_name, section, generated)


def _layout_timeline_table(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Render the 5-column Arabic timeline table for Doc 3 / Doc 4.

    Pulls structured rows from ``generated.allocation.table_rows_ar``
    rather than parsing them back out of the MISSION_TIMELINE
    display strings. Also emits the total / planning / subordinate
    summary lines above the table.
    """
    alloc = generated.allocation
    if alloc is None:
        # Fallback: degrade to numbered_fields when no allocation is
        # available (e.g. a misconfigured template).
        _layout_numbered_fields(ctx, schema_name, section, generated)
        return

    pydantic_instance = generated.sections[schema_name]

    # Emit non-table context fields above the table (current_date,
    # mission_start_time, total_available_time, allocated_planning_time,
    # available_time_for_subordinate_units).
    summary_fields = (
        ("current_date", "التاريخ الحالي"),
        ("mission_start_time", "وقت بدء المهمة"),
        ("total_available_time", "إجمالي الوقت المتاح"),
        ("allocated_planning_time", "الوقت المخصص للتخطيط (القاعدة 1:3)"),
        ("available_time_for_subordinate_units", "الوقت المتاح للوحدات الفرعية"),
    )
    for fname, label in summary_fields:
        if hasattr(pydantic_instance, fname):
            value = getattr(pydantic_instance, fname)
            add_level_one(ctx, f"{label}: {value}")

    # Then the 5-column step table.
    add_paragraph(
        ctx.document,
        "جدول مراحل التخطيط (قاعدة 30/20/30/20):",
        bold=True,
        underline=True,
        rtl=True,
        space_before=0.5,
        space_after=0.0,
    )
    headers = ["النشاط", "النسبة", "المدة", "البدء", "الانتهاء"]
    rows = list(alloc.table_rows_ar)
    add_table(ctx.document, rows, headers, rows_length=5)


# =============================================================================
# §C24 — Y-schema nested layouts
# =============================================================================
# Three layouts below mirror the paragraph hierarchy + timeline-table of the
# OLD generator at ``/Desktop/ToTransfer/New Text Document.txt §6`` — same
# level-1 / level-2 / level-3 / level-4 indentation and the same 5-column
# step table, against the new Y-approved flat Pydantic schemas. Headings are
# military-flavour (the old file was health-themed cover text; the field
# names are identical).
#
# Shared helper: renders the time-allocation block (5 level-twos + the
# 5-column step table) for Doc 3 and Doc 4 without duplicating code.

def _render_y_time_allocation_block(
    ctx: ArabicDocumentContext, instance, allocation
) -> None:
    """Level 1: الإطار الزمني للمهمة, then 5 level_two time fields, then
    the step-allocation table. Matches old doc 3 §1–§1.و + the table.
    """
    add_level_one(ctx, "الإطار الزمني للمهمة.", underline=True, space_before=1.0)

    time_now = getattr(instance, "time_now", "")
    time_Y = getattr(instance, "time_Y", "")
    add_level_two(ctx, f"الوقت الحالي. {time_now} {time_Y}", underline=False)

    add_level_two(
        ctx, f"وقت بدء المهمة (H-Hour). {getattr(instance, 'mission_start', '')}",
        underline=False,
    )
    add_level_two(
        ctx, f"إجمالي الوقت المتاح. {getattr(instance, 'total_available_time', '')}",
        underline=False,
    )
    add_level_two(
        ctx, f"الوقت المخصص للتخطيط. {getattr(instance, 'allocated_planning_time', '')}",
        underline=False,
    )
    add_level_two(
        ctx,
        "الوقت المتاح للوحدات التابعة. "
        f"{getattr(instance, 'available_time_for_subordinate_units', '')}",
        underline=False,
    )
    add_level_two(ctx, "توزيع وقت التخطيط.", underline=True, space_after=0.5)

    headers = ["النشاط", "النسبة المخصصة", "المدة (ساعة)", "وقت البدء", "وقت الانتهاء"]

    if allocation is not None and allocation.table_rows_ar:
        # The four MDMP step rows already carry (name, %, duration, start, end).
        rows = list(allocation.table_rows_ar)
        # Append the "الإجمالي" summary row mirroring old doc 3 layout.
        total_row = (
            "الإجمالي",
            "100%",
            getattr(instance, "allocated_planning_time", ""),
            rows[0][3],  # first step start = planning window start
            rows[-1][4],  # last step end   = planning window end
        )
        rows.append(total_row)
    else:
        # Degrade gracefully without an allocation (shouldn't happen for
        # time_analysis / initial_planning_guidance — their YAMLs compute
        # the allocation via time_math.compute_allocation).
        rows = [
            ("استلام وتحليل المهمة", "30%",
             getattr(instance, "time_for_mission_receipt", ""), "—", "—"),
            ("تطوير الأعمال الممكنة", "20%",
             getattr(instance, "time_for_development", ""), "—", "—"),
            ("تحليل ومقارنة الأعمال", "30%",
             getattr(instance, "time_for_mission_analysis", ""), "—", "—"),
            ("إعداد الخطة والأوامر", "20%",
             getattr(instance, "time_for_plan", ""), "—", "—"),
            ("الإجمالي", "100%",
             getattr(instance, "allocated_planning_time", ""), "—", "—"),
        ]
    add_table(ctx.document, rows, headers, rows_length=len(rows))


def _layout_y_time_analysis(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Y-schema ``time_analysis`` — exact old doc 3 hierarchy.

    One level-1 group, five level-2 time rows, one level-2 sub-heading,
    then the 5-column step table. Field values come from the flat
    ``TimeAnalysis`` Pydantic instance; table rows come from
    ``generated.allocation.table_rows_ar``.
    """
    _reset_level_one_counter(ctx)
    instance = generated.sections[schema_name]
    _render_y_time_allocation_block(ctx, instance, generated.allocation)


def _layout_y_initial_planning_guidance(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Y-schema ``initial_planning_guidance`` — time block (doc 3 style)
    followed by 8 planning-directive level-1 items (doc 4 style).

    Directive paragraphs come from doctrine-retrieval or the Warning-Order
    extractor; each is rendered as a level-1 heading with the value
    inlined. This mirrors old doc 4 §2–§8 exactly.
    """
    _reset_level_one_counter(ctx)
    instance = generated.sections[schema_name]

    _render_y_time_allocation_block(ctx, instance, generated.allocation)

    # §2..§9 — planning directives.
    directive_fields = (
        ("report_production",          "كيفية إنتاج التقارير"),
        ("coordination_duties",        "إجراءات التنسيق"),
        ("authorized_movements",       "الحركات المأذون بها"),
        ("staff_duties",               "واجبات الأركان"),
        ("times_locations_planning",   "أوقات ومواقع التخطيط التشاركي"),
        ("commander_intel_req",        "متطلبات القائد الحرجة (CCIR / PIR)"),
        ("commander_intel_req2",       "متطلبات القوات الصديقة (FFIR)"),
        ("ROE",                        "قواعد الاشتباك (ROE)"),
    )
    for fname, label in directive_fields:
        value = (getattr(instance, fname, "") or "").strip()
        add_level_one(ctx, f"{label}.", underline=True, space_before=1.0)
        if value:
            append_to_paragraph(ctx.document, f" {value}")


def _layout_y_staff_brief(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Y-schema ``staff_brief`` — nested level 1/2/3/4 hierarchy mirroring
    old doc 2 (Epidemiological Brief) structurally.

    Five major sub-sections, each resetting the level-1 counter:
      A. تقدير الاستخبارات والبيئة
      B. تقدير العمليات
      C. تقدير الأفراد
      D. التقدير اللوجستي
      E. الاستنتاجات العملياتية
    """
    instance = generated.sections[schema_name]
    doc = ctx.document

    def lvl1(label: str, value: str | None = None, *, underline: bool = True,
             space_before: float = 1.0) -> None:
        add_level_one(ctx, f"{label}.", underline=underline, space_before=space_before)
        if value:
            append_to_paragraph(doc, f" {value}")

    def lvl2(label: str, value: str | None = None, *, underline: bool = True) -> None:
        add_level_two(ctx, f"{label}.", underline=underline)
        if value:
            append_to_paragraph(doc, f" {value}")

    def lvl3(label: str, value: str | None = None, *, underline: bool = True) -> None:
        add_level_three(ctx, f"{label}.", underline=underline)
        if value:
            append_to_paragraph(doc, f" {value}")

    def lvl4(label: str, value: str | None = None, *, underline: bool = True) -> None:
        add_level_four(ctx, f"{label}.", underline=underline)
        if value:
            append_to_paragraph(doc, f" {value}")

    def section_header(text: str) -> None:
        _reset_level_one_counter(ctx)
        add_paragraph(
            doc, text, rtl=True, underline=True, bold=True,
            space_before=1.0, space_after=0.3,
        )

    # =============== A. تقدير الاستخبارات والبيئة ===============
    section_header("تقدير الاستخبارات والبيئة")

    lvl1("الأرض", instance.Terrain)
    lvl1("الطقس")
    lvl2("الحالة العامة", instance.Weather)
    lvl2("التنبؤات الجوية", underline=True)
    lvl3("أول ضوء (BMNT)", instance.First_light)
    lvl3("آخر ضوء (EENT)", instance.Last_light)
    lvl3("طور القمر", instance.Moon)
    lvl1(
        "تأثير الأرض والطقس على العمليات",
        instance.Effect_of_Weather_and_Terrain_on_Operations,
    )

    lvl1("موقف العدو")
    lvl2("تشكيل العدو", instance.Composition)
    lvl2("انتشار العدو")
    lvl3("التفاصيل", instance.Deployments)
    lvl2("تغطية قوات العدو", instance.Force_Coverage)
    lvl2("معنويات العدو", instance.Morale)
    lvl2("تدريب العدو", instance.Training)
    lvl2("النشاطات الأخيرة والجارية")
    lvl3("التفاصيل", instance.Recent_and_Ongoing_Activities)
    lvl2("تكتيكات العدو في مراحل العملية")
    lvl3("المرحلة الأولى — التحضير")
    lvl4("الأسلوب", instance.Enemy_Tactics_in_Exposure_Operations_Phase1_Preparation)
    lvl3("المرحلة الثانية — التشكيل")
    lvl4("الأسلوب", instance.Enemy_Tactics_in_Exposure_Operations_Phase2_Preparation)
    lvl3("المرحلة الثالثة — الهجوم الرئيسي")
    lvl4("الأسلوب", instance.Enemy_Tactics_in_Exposure_Operations_Phase3_Main_Attack)
    lvl2("نوايا العدو وأهدافه", instance.Intentions_and_Objectives)
    lvl2("ملاحظات الاستخبارات المضادة")
    lvl3("التفاصيل", instance.Counter_Intelligence_Observations)
    lvl2("قدرات العدو")
    lvl3("التفاصيل", instance.Enemy_Capabilities)

    lvl1("الاستنتاجات الاستخباراتية")
    lvl2("التفاصيل", instance.Conclusions)

    # =============== B. تقدير العمليات ===============
    section_header("تقدير العمليات")

    lvl1("مهمة القيادة الأعلى")
    lvl2("القيادة المشتركة", underline=True)
    lvl3("المهمة", instance.join_op_mission)
    lvl3("الرؤية", underline=True)
    lvl4("الغاية", instance.Join_op_purp)
    lvl4("النسق", instance.joint_ops_how)
    lvl4("الحالة النهائية المرغوبة", instance.joint_ops_desired_end)
    lvl2("القيادة التنفيذية", underline=True)
    lvl3("المهمة", instance.Exc_command_mission)
    lvl3("الرؤية", underline=True)
    lvl4("الغاية", instance.Exc_command_purp)
    lvl4("النسق", instance.Exc_command_main_mission)
    lvl4("الحالة النهائية المرغوبة", instance.joint_ops_desired_end2)

    lvl1("القوة البرية", instance.Land_component_force)
    lvl1("الوحدات المرتبطة", instance.Attached_units)
    lvl1("تشكيل القوة الصديقة")
    lvl2("التأليف", instance.Force_Composition)
    lvl2("مستوى الجاهزية التدريبية", instance.Training_Readiness_Level)
    lvl2("الفاعلية القتالية", instance.Combat_Effectiveness)

    # =============== C. تقدير الأفراد ===============
    section_header("تقدير الأفراد")

    lvl1("تغطية القوات الصديقة", instance.Force_Cover)
    lvl1("المعنويات القتالية", instance.Combat_Morale)
    lvl1("التعزيزات", instance.Reinforcements)
    lvl1("الخسائر المتوقَّعة", instance.Projected_Casualties)
    lvl1("التحكم والتنسيق", instance.Control_and_Coordination)
    lvl1("أسرى الحرب", instance.Prisoners_of_War)
    lvl1("المدنيون في منطقة العمليات", instance.Civilian_Users)
    lvl1("المحتجزون والمعتقلون المدنيون", instance.Civilian_Prisoners_and_Detainees)
    lvl1("استنتاجات تقدير الأفراد")
    lvl2("التفاصيل", instance.Human_Force_Conclusions)

    # =============== D. التقدير اللوجستي ===============
    section_header("التقدير اللوجستي")

    lvl1("الإعاشة", instance.Logistical_Rations)
    lvl1("الإدامة العامة", instance.Logistical_sustainment)
    lvl1("الوقود", instance.Fuel)
    lvl1("الذخيرة", instance.ammunition)
    lvl1("قطع الغيار", instance.Spare_parts)
    lvl1("العتاد الهندسي", instance.Engineering_materiel)
    lvl1("النقل", instance.Transportation)
    lvl1("الصيانة", instance.Maintenance)
    lvl1("المستشفيات الميدانية", instance.Field_Hospitals)
    lvl1("استنتاجات التقدير اللوجستي")
    lvl2("التفاصيل", instance.Supply_Conclusions)

    # =============== E. الاستنتاجات العملياتية ===============
    section_header("الاستنتاجات العملياتية")

    lvl1("التفاصيل", instance.Operational_Conclusions)


def _layout_y_warning_order(
    ctx: ArabicDocumentContext, schema_name: str, section, generated
) -> None:
    """Y-schema ``warning_order`` — mirrors old generator doc 1.

    Source of truth: ``/Desktop/ToTransfer/New Text Document.txt`` lines
    939–1152 (the ``document = Document()`` block — 40/50 Y fields
    rendered 1:1 with old-doc-1 hierarchy, military Arabic labels
    replacing the health-themed cover text). The remaining 10 Y fields
    (``header4``, ``situation``, ``mission_of_supporting_unit``,
    ``join_op_{mission,purp,ops_how,ops_desired_end}``, ``date_time``,
    ``local_authorities``, ``red_crescent``) are inserted at the
    nearest doctrinally-sensible anchor so no Y field is silently
    dropped (the dispatcher's no-empty-strings post-condition still
    holds).

    Layout hierarchy (doc 1 verbatim):

      * "بسم الله الرحمن الرحيم" centred
      * ``add_arabic_header`` block (organization/main_unit/sub_unit,
        place = Assembly_Area, today Hijri + Gregorian, letter_ref_number)
      * letter_ref_number2 centred underlined
      * References / Maps / time_zone / task_assembly as plain paragraphs
      * LEVEL 1 "الموقف" → area_interest / operations_area (level 2);
        terrain / weather / civil_considerations (level 3); enemy_forces
        (level 2); friendly_forces (level 2 + SPLITTER); gov_and_nongov_org
        with local_authorities + red_crescent sub-items; CIVILIAN_CONSIDERATIONS;
        Attached_and_Detached_units (SPLITTER); Operational_Assumptions (SPLITTER).
      * LEVEL 1 "مهمة المكون البري" → GROUND_COMPONENT_MISSION +
        join-op block (mission/purp/how/desired-end) + mission_of_supporting_unit.
      * LEVEL 1 "التنفيذ" → Exc_command_purp / Concept_of_operations (level 2);
        Units_Duty / Duties_of_Other (SPLITTER); LEVEL-2 "تعليمات التنسيق"
        with Timings / CCIR / Fire / Air / Risk / ROE / Media / Meeting /
        Excu / Movm as LEVEL-3 children.
      * LEVEL 1 "الإدامة" + Sustainment (SPLITTER).
      * LEVEL 1 "القيادة والسيطرة" + ACCS.
      * "أقرّوا:" approval block (3 signature lines).
      * "الملاحق:" → Appendices split by line via add_level_one_ML.
      * "الشفافات:" → Viewports split by line via add_level_one_SHFAF.
    """
    from graph.generation.time_math import (
        format_gregorian_date,
        format_hijri_date,
        gregorian_to_hijri,
    )

    instance = generated.sections[schema_name]
    doc = ctx.document

    def _s(val: Any) -> str:
        """Stringify + trim; treat None as empty."""
        return str(val).strip() if val is not None else ""

    def lvl1(label: str, value: str | None = None, *, underline: bool = True,
             space_before: float = 0.0) -> None:
        add_level_one(ctx, f"{label}.", underline=underline, space_before=space_before)
        v = _s(value)
        if v:
            append_to_paragraph(doc, f" {v}")

    def lvl2(label: str, value: str | None = None, *, underline: bool = True) -> None:
        add_level_two(ctx, f"{label}.", underline=underline)
        v = _s(value)
        if v:
            append_to_paragraph(doc, f" {v}")

    def lvl3(label: str, value: str | None = None, *, underline: bool = True) -> None:
        add_level_three(ctx, f"{label}.", underline=underline)
        v = _s(value)
        if v:
            append_to_paragraph(doc, f" {v}")

    # ---------------------------------------------------- 1. bism + header block
    _reset_level_one_counter(ctx)

    add_paragraph(
        doc, "بسم الله الرحمن الرحيم",
        rtl=True, middle_alignment=True,
        bold=False, underline=False, space_after=1.0,
    )

    today = datetime.now()
    date_greg = format_gregorian_date(today)
    date_hijri = format_hijri_date(gregorian_to_hijri(today))

    # `header` field (primary classification / copy-number). Fall back
    # to a blank "نسخة رقم ... من ..." line so the header block still
    # renders when the source file didn't carry this line.
    copy_number = _s(instance.header) or "نسخة رقم (  ) من (  )"

    add_arabic_header(
        doc=doc,
        copy_number_placeholder=copy_number,
        organization=_s(instance.header2),
        main_unit=_s(instance.header3),
        sub_unit=_s(instance.header4),
        place=_s(instance.Assembly_Area),
        date_hijri=date_hijri,
        date_greg=date_greg,
        letter_ref_number=f"الرقم: {_s(instance.letter_ref_number)}",
        target_length_chars=60,
        max_kashidas_per_word=30,
        cell_width_right=Cm(15),
        cell_width_left=Cm(1),
    )

    # Reference number 2 (centred, underlined — doc 1 line 969).
    add_paragraph(
        doc, _s(instance.letter_ref_number2),
        rtl=True, middle_alignment=True,
        bold=False, underline=True, space_before=1.0,
    )

    # Scenario-declared date/time line — Y-specific (not in doc 1).
    # Kept as a plain paragraph under letter_ref_number2 so the live
    # date in add_arabic_header stays "today" (matching doc 1) while
    # the operation's own date is preserved in the rendering.
    if _s(instance.date_time):
        add_paragraph(
            doc, f"التاريخ والوقت. {_s(instance.date_time)}",
            rtl=True, bold=False, space_before=0.3,
        )

    # References / Maps / time_zone / task_assembly — doc 1 lines 972–982.
    add_paragraph(
        doc, f"المراجع: {_s(instance.References)}",
        rtl=True, bold=False, space_before=1.0,
    )
    add_paragraph(
        doc, f"الخرائط. {add_full_stop(_s(instance.Maps))}",
        rtl=True,
    )
    add_paragraph(doc, _s(instance.time_zone), rtl=True)
    add_paragraph(doc, _s(instance.task_assembly), rtl=True)

    # ---------------------------------------------------- 2. الموقف (Situation)
    lvl1("الموقف", space_before=1.0)

    # Y field `situation` — short preamble paragraph under level 1.
    if _s(instance.situation):
        append_to_paragraph(doc, f" {_s(instance.situation)}")

    lvl2("منطقة الاهتمام", add_full_stop(_s(instance.area_interest)))
    lvl2("منطقة العمليات", add_full_stop(_s(instance.operations_area)))
    lvl3("الأرض", instance.terrain)
    lvl3("الطقس", instance.weather)
    lvl3("الاعتبارات المدنية في منطقة العمليات", instance.civil_considerations)
    lvl2("قوات العدو", instance.enemy_forces)

    lvl2("القوات الصديقة", underline=True)
    SPLITTER(_s(instance.friendly_forces), ctx)

    lvl2("الجهات الحكومية والمنظمات غير الحكومية", instance.gov_and_nongov_org)
    lvl3("السلطات المحلية", instance.local_authorities)
    lvl3("الهلال الأحمر", instance.red_crescent)

    lvl2("الاعتبارات المدنية", instance.CIVILIAN_CONSIDERATIONS)

    lvl2("الوحدات المرتبطة والمنفصلة", underline=True)
    SPLITTER(_s(instance.Attached_and_Detached_units), ctx)

    lvl2("الافتراضات العملياتية", underline=True)
    SPLITTER(_s(instance.Operational_Assumptions), ctx)

    # ---------------------------------------------------- 3. مهمة المكون البري
    lvl1("مهمة المكون البري", space_before=1.0)
    append_to_paragraph(doc, f" {_s(instance.GROUND_COMPONENT_MISSION)}")

    # Y-additional: higher-command (joint-op) framing — nested level-2
    # to distinguish higher-command intent from our own mission line.
    lvl2("مهمة القيادة المشتركة", instance.join_op_mission)
    lvl3("الغاية", instance.join_op_purp)
    lvl3("النسق", instance.joint_ops_how)
    lvl3("الحالة النهائية المرغوبة", instance.joint_ops_desired_end)
    lvl2("مهمة الوحدات المساندة", instance.mission_of_supporting_unit)

    # ---------------------------------------------------- 4. التنفيذ (Execution)
    lvl1("التنفيذ", space_before=1.0)

    lvl2("غاية القائد", instance.Exc_command_purp)
    lvl2("مفهوم العمليات", instance.Concept_of_operations)

    lvl2("واجبات الوحدات التابعة", underline=True)
    SPLITTER(_s(instance.Units_Duty), ctx)

    lvl2("واجبات وحدات القتال والإسناد الأخرى", underline=True)
    SPLITTER(_s(instance.Duties_of_Other_Combat_Units_and_Combat_Support_Units), ctx)

    lvl2("تعليمات التنسيق", underline=True)
    lvl3("التوقيتات", instance.Timings)

    lvl3("متطلبات القائد الحرجة من المعلومات (CCIR)", underline=True)
    SPLITTER(_s(instance.Commanders_Crtitical_Information_Requirements), ctx)

    lvl3("تنسيق الإسناد الناري", instance.Fire_support_coordination)
    lvl3("تنسيق الإسناد الجوي", instance.Air_support_coordination)
    lvl3("تقدير المخاطر", instance.Risk_assy)

    lvl3("قواعد الاشتباك (ROE)", underline=True)
    SPLITTER(_s(instance.ROE), ctx)

    lvl3("الإعلام والمعلومات العامة", instance.Other_coordination_media)
    lvl3("اجتماعات التنسيق", instance.Other_coordination_meeting)

    lvl3("تعليمات التنفيذ الفورية", underline=True)
    SPLITTER(_s(instance.Other_coordination_Excu), ctx)

    lvl3("تنسيق الحركة", instance.Other_coordination_movm)

    # ---------------------------------------------------- 5. الإدامة
    lvl1("الإدامة", space_before=1.0)
    SPLITTER(_s(instance.Sustainment), ctx)

    # ---------------------------------------------------- 6. القيادة والسيطرة
    lvl1("القيادة والسيطرة", space_before=1.0)
    append_to_paragraph(doc, f" {_s(instance.ACCS)}")

    # ---------------------------------------------------- 7. Approval block
    add_paragraph(doc, "أقرّوا:", rtl=True, underline=False, space_before=1.0)
    add_paragraph(doc, "عن / قائد التشكيل", rtl=True, tabs=7, underline=False)
    add_paragraph(doc, "قائد العمليات", rtl=True, tabs=7, underline=False)
    add_paragraph(doc, "رئيس هيئة الأركان", rtl=True, tabs=7, underline=False)

    # ---------------------------------------------------- 8. Appendices
    add_paragraph(doc, "الملاحق:", rtl=True, underline=False, space_before=3.0)
    for line in _s(instance.Appendices).splitlines():
        if line.strip():
            clean = line.strip().replace("\t", " ")
            add_level_one_ML(ctx, add_full_stop(clean))

    # ---------------------------------------------------- 9. Viewports (slides)
    add_paragraph(doc, "الشفافات:", rtl=True, underline=False, space_before=1.0)
    for line in _s(instance.Viewports).splitlines():
        if line.strip():
            clean = line.strip().replace("\t", " ")
            add_level_one_SHFAF(ctx, add_full_stop(clean))


_LAYOUT_RENDERERS = {
    "numbered_fields": _layout_numbered_fields,
    "timeline_table": _layout_timeline_table,
    "staff_sections": _layout_staff_sections,
    # §C24 — Y-schema nested layouts mirroring old generator §6 hierarchy.
    "y_time_analysis": _layout_y_time_analysis,
    "y_initial_planning_guidance": _layout_y_initial_planning_guidance,
    "y_staff_brief": _layout_y_staff_brief,
    "y_warning_order": _layout_y_warning_order,
}


def _render_section(
    ctx: ArabicDocumentContext, entry, generated
) -> None:
    schema_name = _extra_get(entry, "schema")
    heading = _extra_get(entry, "heading")
    layout = _extra_get(entry, "layout") or "numbered_fields"

    _render_heading(ctx, heading)

    renderer = _LAYOUT_RENDERERS.get(layout)
    if renderer is None:
        # Unknown layouts degrade to numbered_fields; M3+ will register
        # the proper hooks for header_block / staff_sections / etc.
        renderer = _layout_numbered_fields
    renderer(ctx, schema_name, entry, generated)


def _render_approval_block(ctx: ArabicDocumentContext, entry) -> None:
    lines = _extra_get(entry, "lines") or []
    add_paragraph(
        ctx.document,
        "أقرّوا:",
        bold=True,
        underline=True,
        rtl=True,
        space_before=1.0,
        space_after=0.3,
    )
    for line in lines:
        add_paragraph(
            ctx.document,
            str(line),
            bold=False,
            underline=False,
            rtl=True,
            space_before=0.2,
        )


def _render_appendices_list(ctx: ArabicDocumentContext, entry, generated) -> None:
    source = _extra_get(entry, "source")
    if not source:
        return
    # `source` names a field on some schema (usually Annexes.appendices).
    for _schema_name, sec in generated.sections.items():
        if hasattr(sec, source):
            value = str(getattr(sec, source))
            add_paragraph(
                ctx.document,
                "الملاحق:",
                bold=True,
                underline=True,
                rtl=True,
                space_before=0.5,
            )
            add_paragraph(ctx.document, value, rtl=True)
            return


def _render_page_break(ctx: ArabicDocumentContext) -> None:
    ctx.document.add_page_break()


# =============================================================================
# 9. CITATION ENDNOTES (new — ref 19 §6)
# =============================================================================

@dataclass(frozen=True)
class CitationEntry:
    """One row of the citation-endnote block.

    Sequential ``number`` is assigned in walk order; ``full_tag`` is
    the exact ``[source_doc §locator]`` (legacy) or ``[S/O/D: <slug>
    §<locator>]`` (tier-aware, §C31) string that appears inline in
    the drafted prose. The endnote reads
    ``[N] source_doc — فقرة locator`` per ref 19 §6.

    ``tier`` (added §C31, Phase 6) labels which evidence channel the
    entry came from: ``"source_files"`` / ``"operationalfiles"`` /
    ``"doctrine"`` / ``"legacy"``. ``"legacy"`` covers entries
    produced by groups that resolved via the pre-§C32 ``collections:``
    YAML key — they emit the untagged citation shape and render as
    a flat list under the original ``الاستشهادات`` heading. Mixed
    templates render the tier-aware entries under three Arabic
    sub-headings + the legacy ones under a fallback heading.
    """

    number: int
    source_doc: str
    locator: str
    full_tag: str
    collection: str  # originating Qdrant collection (SourcedHit provenance)
    tier: str = "legacy"


def collect_citations(generated) -> list[CitationEntry]:
    """Walk the bundle's three channels (source_files / operationalfiles
    / doctrine) when present; otherwise fall back to ``retrieval_results``.

    Duplicates are folded into one entry — the drafter might cite
    the same chunk from two different fields; we want one endnote
    row per source, not per inline occurrence.

    Returns an empty list when no retrieved fields exist (Doc 3)
    or when every group's hits list is empty.
    """
    entries: list[CitationEntry] = []
    seen: set[str] = set()

    bundles = getattr(generated, "evidence_bundles", ()) or ()

    if bundles:
        # §C31 path — bundle is the source of truth so the renderer
        # can include FactSnippets (source_files) alongside chunk
        # hits.  Order: source_files → operationalfiles → doctrine
        # within each group, in group declaration order.
        for bundle in bundles:
            for snip in getattr(bundle, "source_files_evidence", ()) or ():
                kind = snip.source_file_kind or snip.field_name or "source_file"
                tag = f"[S: {kind} §extracted]"
                if tag in seen:
                    continue
                seen.add(tag)
                entries.append(CitationEntry(
                    number=len(entries) + 1,
                    source_doc=kind,
                    locator="extracted",
                    full_tag=tag,
                    collection="",  # no Qdrant origin for FactSnippet
                    tier="source_files",
                ))
            for sh in getattr(bundle, "operationalfiles_evidence", ()) or ():
                _maybe_append_hit_entry(entries, seen, sh, default_tier="operationalfiles")
            for sh in getattr(bundle, "doctrine_evidence", ()) or ():
                _maybe_append_hit_entry(entries, seen, sh, default_tier="doctrine")
        return entries

    # Legacy path — pre-§C29 callers (or templates that don't go
    # through the dispatcher's bundle assembly) walk retrieval_results
    # directly.  Each SourcedHit carries its own ``tier`` field so we
    # honour it here too in case Phase 7 wires non-bundle paths later.
    for retrieval in getattr(generated, "retrieval_results", ()):
        for sh in retrieval.hits:
            _maybe_append_hit_entry(entries, seen, sh, default_tier="legacy")
    return entries


def _maybe_append_hit_entry(
    entries: list[CitationEntry],
    seen: set[str],
    sh,
    *,
    default_tier: str,
) -> None:
    """Helper: dedupe by full_tag and append a CitationEntry."""
    tag = sh.citation_tag
    if not tag or tag in seen:
        return
    seen.add(tag)
    parsed = _parse_citation_tag(tag, sh.hit.source_doc)
    slug, locator, parsed_tier = parsed
    # Prefer the tier carried on the SourcedHit (Phase 7+ source of
    # truth); fall back to the parsed prefix; final fallback is the
    # caller-provided default.
    sh_tier = getattr(sh, "tier", None)
    tier = sh_tier or parsed_tier or default_tier
    entries.append(CitationEntry(
        number=len(entries) + 1,
        source_doc=slug,
        locator=locator,
        full_tag=tag,
        collection=getattr(sh, "collection", ""),
        tier=tier,
    ))


def _parse_citation_tag(tag: str, source_doc_fallback: str) -> tuple[str, str, str]:
    """Split a citation tag into ``(slug, locator, tier)``.

    Recognises three shapes:
      * Legacy:        ``[<slug> §<locator>]``                   → tier=""
      * Tier-aware:    ``[O: <slug> §<locator>]``                → tier="operationalfiles"
                       ``[D: <slug> §<locator>]``                → tier="doctrine"
                       ``[S: <kind> §<locator>]``                → tier="source_files"

    Falls back gracefully: a non-conforming tag becomes
    ``(source_doc_fallback minus .pdf, raw inner, "")`` so the
    endnote row still renders.
    """
    inner = tag.strip()
    if inner.startswith("[") and inner.endswith("]"):
        inner = inner[1:-1]

    # Tier prefix detection — single letter + colon + space.
    tier = ""
    if inner[:3] in ("S: ", "O: ", "D: "):
        prefix_letter = inner[0]
        tier = {"S": "source_files", "O": "operationalfiles", "D": "doctrine"}[prefix_letter]
        inner = inner[3:]

    if " §" in inner:
        slug, _, locator = inner.partition(" §")
        return slug.strip(), locator.strip(), tier

    # Fallback: source_doc as slug, whole tag (post-prefix-strip) as locator.
    fb = source_doc_fallback
    if fb.lower().endswith(".pdf"):
        fb = fb[:-4]
    return fb, inner, tier


_TIER_SUB_HEADINGS_AR = {
    "source_files": "ملفات مرفوعة من المستخدم",
    "operationalfiles": "المصادر التشغيلية",
    "doctrine": "المرجع العقيدي",
}


def render_citations_section(ctx: ArabicDocumentContext, entries: list[CitationEntry]) -> None:
    """Append the ``الاستشهادات`` endnote block to the document.

    Heading is a level-0 titled paragraph (bold + underline, like
    every other top-level section heading in the port). Each entry
    is emitted as a right-aligned paragraph ``[N] slug — فقرة locator``
    with the usual Arial 14 pt / 0.75 cm exact line spacing.

    Layout (§C31, Phase 6):
      * Pure-legacy entries (``tier="legacy"`` for every entry) →
        flat list under the original ``الاستشهادات`` heading; byte-
        equal to the pre-§C31 output for every existing template.
      * Tier-aware entries → group by tier under three Arabic sub-
        headings (``ملفات مرفوعة من المستخدم`` /
        ``المصادر التشغيلية`` / ``المرجع العقيدي``); empty tiers are
        omitted entirely. Mixed templates render legacy entries last
        under ``مصادر`` so their fallback shape is still visible
        without losing the sub-heading structure for the rest.
    """
    if not entries:
        return

    add_paragraph(
        ctx.document,
        "الاستشهادات",
        bold=True,
        underline=True,
        rtl=True,
        space_before=1.0,
        space_after=0.3,
    )

    # Detect whether the layout should switch to per-tier sub-headings.
    # Sub-headings exist to disambiguate sources across channels; with a
    # single populated tier (e.g. legacy template that only ever hits
    # operationalfiles, or a tier-aware template whose policy is
    # ``doctrine_only``) there is nothing to disambiguate, so the flat
    # layout is byte-equal-correct.  Trigger sub-headings only when 2+
    # named tiers carry entries.  This keeps pre-§C31 layout behaviour
    # for legacy templates even after §C29 added the default
    # ``tier="operationalfiles"`` to ``SourcedHit``.
    populated_tiers = {e.tier for e in entries if e.tier in _TIER_SUB_HEADINGS_AR}
    has_tiered = len(populated_tiers) > 1

    if not has_tiered:
        for entry in entries:
            add_paragraph(
                ctx.document,
                f"[{entry.number}] {entry.source_doc} — فقرة {entry.locator}",
                bold=False,
                underline=False,
                rtl=True,
                space_before=0.1,
                space_after=0.0,
            )
        return

    # Tiered layout — group by tier in canonical order.  Empty tiers
    # are skipped; legacy entries fall under a "مصادر" catch-all so
    # mixed templates don't lose their flat-shape rows.
    for tier_key in ("source_files", "operationalfiles", "doctrine"):
        tier_entries = [e for e in entries if e.tier == tier_key]
        if not tier_entries:
            continue
        add_paragraph(
            ctx.document,
            _TIER_SUB_HEADINGS_AR[tier_key],
            bold=True,
            underline=False,
            rtl=True,
            space_before=0.4,
            space_after=0.1,
        )
        for entry in tier_entries:
            add_paragraph(
                ctx.document,
                f"[{entry.number}] {entry.source_doc} — فقرة {entry.locator}",
                bold=False,
                underline=False,
                rtl=True,
                space_before=0.1,
                space_after=0.0,
            )

    legacy_entries = [e for e in entries if e.tier == "legacy"]
    if legacy_entries:
        add_paragraph(
            ctx.document,
            "مصادر",
            bold=True,
            underline=False,
            rtl=True,
            space_before=0.4,
            space_after=0.1,
        )
        for entry in legacy_entries:
            add_paragraph(
                ctx.document,
                f"[{entry.number}] {entry.source_doc} — فقرة {entry.locator}",
                bold=False,
                underline=False,
                rtl=True,
                space_before=0.1,
                space_after=0.0,
            )


def render_document(generated, output_path: Path | str) -> Path:
    """Render a :class:`GeneratedDocument` to a ``.docx`` at ``output_path``.

    Walks ``generated.template.structure`` once, dispatching on each
    entry's ``kind`` (``title``, ``section``, ``approval_block``,
    ``appendices_list``, ``page_break``). Returns the resolved path
    for convenience.
    """
    ctx = ArabicDocumentContext.new()
    configure_document(ctx.document)

    for entry in generated.template.structure:
        kind = entry.kind
        if kind == "title":
            _render_title(ctx, entry)
        elif kind == "section":
            _render_section(ctx, entry, generated)
        elif kind == "approval_block":
            _render_approval_block(ctx, entry)
        elif kind == "appendices_list":
            _render_appendices_list(ctx, entry, generated)
        elif kind == "page_break":
            _render_page_break(ctx)
        else:
            # Unknown structure kinds fall through silently — M4+
            # may register additional renderers.
            continue

    # Implicit citation endnotes (scoping §16 D7 / ref 19 §6).
    # Appended at "document end" as the scoping doc prescribes,
    # without requiring every template author to remember to declare
    # an explicit `citations` structure entry. No-ops for Doc 3
    # (retrieval_results is empty).
    citations = collect_citations(generated)
    if citations:
        render_citations_section(ctx, citations)

    configure_last_page_section(ctx.document)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    ctx.document.save(str(out))
    return out
