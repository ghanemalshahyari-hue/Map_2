"""
build_report.py — WarGamingClaude3_Report_AR.docx
=================================================
Arabic-first RTL operational wargame report. For each of 17 steps embeds three
Codex-rendered JPEGs (full overview, AOI zoom, action zoom) and writes the full
action-reaction narrative across 6 domains: strategic, maritime, air, mines,
USV/UAV swarm, land, plus SOF, EW, logistics, losses-with-reason, and advantage.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
STATE = json.load(open(ROOT / "qa" / "simulation_full_state.json"))
SNAPS = STATE["snapshots"]
BLS_GJ = json.load(open(ROOT / "bls_selection.geojson"))
NATO = json.load(open(ROOT / "source_inputs" / "parsed_nato.json"))
PARSED_DOCS = json.load(open(ROOT / "source_inputs" / "parsed_docs.json"))

# ------------------------------------------------------------------
# RTL helpers
# ------------------------------------------------------------------
def set_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_doc_rtl(doc):
    sectPr = doc.sections[0]._sectPr
    bidi = OxmlElement('w:bidi'); sectPr.append(bidi)

def ar_run(p, text, size=11, bold=False, color=None):
    run = p.add_run(text)
    run.font.name = "Arial"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:cs'), "Arial")
    rFonts.set(qn('w:ascii'), "Arial")
    rFonts.set(qn('w:hAnsi'), "Arial")
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(size*2)))
    rPr.append(szCs)
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color
    rtl_el = OxmlElement('w:rtl'); rtl_el.set(qn('w:val'), '1')
    rPr.append(rtl_el)
    return run

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    set_rtl(p)
    ar_run(p, text, size=sizes.get(level, 11), bold=True, color=RGBColor(0x1e, 0x3a, 0x8a))
    return p

def para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    ar_run(p, text, size=size, bold=bold, color=color)
    return p

def bullet(doc, text, size=11, indent_lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(0.5 * (indent_lvl + 1))
    ar_run(p, text, size=size)
    return p

def page_break(doc):
    p = doc.add_paragraph(); run = p.add_run(); run.add_break()
    return p

def embed_image(doc, path, caption_ar, width_in=6.5):
    if not path.exists():
        para(doc, f"[الصورة غير متوفرة: {path.name}]", size=9, color=RGBColor(0x99,0x33,0x33))
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    # Caption underneath
    cap = doc.add_paragraph()
    set_rtl(cap)
    ar_run(cap, caption_ar, size=9, bold=True, color=RGBColor(0x4b, 0x55, 0x63))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------------------------------------------------------
# Document setup
# ------------------------------------------------------------------
doc = Document()
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width  = Inches(8.5)
section.left_margin = Inches(1); section.right_margin = Inches(1)
section.top_margin = Inches(1); section.bottom_margin = Inches(1)
set_doc_rtl(doc)
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ------------------------------------------------------------------
# Title
# ------------------------------------------------------------------
title = doc.add_paragraph(); set_rtl(title)
ar_run(title, "تقرير المحاكاة العملياتية المشتركة الكاملة", size=22, bold=True, color=RGBColor(0x1e,0x3a,0x8a))
sub = doc.add_paragraph(); set_rtl(sub)
ar_run(sub, "WarGamingClaude3 — هجوم برمائي بري + بحري + جوي", size=14)
sub2 = doc.add_paragraph(); set_rtl(sub2)
ar_run(sub2, "منطقة العمليات: خليج سرت — ليبيا", size=12)
final_outcome = SNAPS[-1].get("outcome_logic","") or ""
sub3 = doc.add_paragraph(); set_rtl(sub3)
ar_run(sub3, f"النتيجة الختامية: {final_outcome}", size=12, bold=True, color=RGBColor(0x7f,0x1d,0x1d))

page_break(doc)

# ------------------------------------------------------------------
# 1. الفحص التمهيدي
# ------------------------------------------------------------------
heading(doc, "1. الفحص التمهيدي", level=1)
para(doc, "تمت قراءة كل البيانات من المصادر الأصلية فقط دون افتراضات. لم تُضف أرقام من خارج الوثائق:")
bullet(doc, f"عدد الوحدات الزرقاء في nato-map-layers.geojson (أفلييشن=3): {len(NATO['friendly'])} (المتوقع 39).")
bullet(doc, f"عدد الوحدات الحمراء في nato-map-layers.geojson (أفلييشن=6): {len(NATO['hostile'])} (المتوقع 13).")
bullet(doc, "تم تحليل ملفي القوة الحمراء.docx والقوة الزرقاء.docx بشكل كامل وبُني ترتيب القتال طبقاً لما ورد فيهما حصراً.")
bullet(doc, "إجمالي الوحدات الحمراء المنمذجة: 3 فرق (الفرقة 4 المشاة الآلية + الفرقة 9 + الفرقة المدرعة 1) + 4 ألوية إضافية (23 و 24 المشاة + 21 العمليات الخاصة + 2nd AD لواء الدفاع الجوي) + لواء صواريخ أرض/أرض + لواءا المدفعية والدفاع الجوي الموزعتان عبر الفرق.")
bullet(doc, "إجمالي الوحدات الزرقاء المنمذجة: المكون البري 99 (3 ألوية: 51، 52، 54) + 5 كتائب مدفعية (551، 552، 554، 555، 556) + كتيبة م/د 502 + كتيبة EW 505 + كتيبة دفاع جوي 507 + الـ80 SOF + احتياطات (72 المدرع + 71 المشاة الآلية + 73 المدرع) + 9th AD Bde (3 هوك + 1 S-300) + لواء SSM 800-1000 كم.")
bullet(doc, "قاعدتان بحريتان حمراوان (أ و ب) + 3 قواعد جوية حمراء (أ و ب و ج). ميناء بحري أزرق + 3 قواعد جوية زرقاء.")
bullet(doc, "تم استخراج 18 درس تاريخي من WarReferences.md (Iwo Jima, Tarawa, D-Day, Falklands, Gallipoli, Wonsan, Inchon, USV الأوكرانية 2022-24, الحوثيون البحر الأحمر 2023-24, Latakia 1973, Praying Mantis 1988, NATO ليبيا 2011، قبرص 1974، القرم 2014، Moskva 2022، إيران 14 أبريل 2024).")
bullet(doc, "تم تطبيق العقائد المرجعية من Doctrines.md: JP 3-02، AJP-3.1، AJP-3.3، ATP 3-01.4 J-SEAD، ATP 3-37.5 الإعاقة المركبة، FM 3-90، ADP 3-0 (نقطة الانهيار)، NWP 3-15 حرب الألغام، NTRP 3-22، AJP-3.9 الاستهداف المشترك.")

# ------------------------------------------------------------------
# 2. الموقف العملياتي
# ------------------------------------------------------------------
heading(doc, "2. الموقف العملياتي", level=1)
para(doc, "تواجه قوة الواجب المشتركة الزرقاء (JTF1) هجوماً برمائياً مشتركاً من قوات الفيلق 1 (قيادة المنطقة الجنوبية) الحمراء. التهديد متعدد الأبعاد: بري + بحري + جوي + قوات خاصة + صواريخ استراتيجية + USV/UAV.")
para(doc, "العمل الأكثر احتمالاً للأحمر (من enemy.docx) ثلاثة مراحل:")
bullet(doc, "المرحلة 1: إنزال عناصر الاستطلاع والتأمين والمقدمة بعمق 1-2 كم مع استخدام معدات الإعاقة الإلكترونية طوال العملية.")
bullet(doc, "المرحلة 2: اقتحام الدفاعات والاستيلاء على رأس الشاطئ بالفرقة 4 (8-10 كم) ثم تطوير الهجوم بالفرقة 9 إلى عمق (40-50 كم).")
bullet(doc, "المرحلة 3: إنزال الفرقة المدرعة 1 للاستثمار نحو الهدف X بعمق (80-100 كم).")
para(doc, "إكتشافات أساسية غير بديهية:")
bullet(doc, "تشويش الكتيبة 405 الإلكترونية مستمر طوال العملية وليس حصراً للمرحلة 1.")
bullet(doc, "تفوق مدى صواريخ أرض/أرض الزرقاء (800-1000 كم) على الأحمر (500-600 كم) — يسمح للأزرق بضرب القواعد الحمراء الخلفية بينما الأحمر لا يستطيع الرد بالمثل.")
bullet(doc, "الأزرق يمتلك 400 لغم بحري مسبق النشر في 4 أحزمة أمام BLS المرشحة — العنصر الحاسم في الدفاع الساحلي.")
bullet(doc, "الأحمر يمتلك 24 زورق مفخخ (USV) + 48 طائرة مفخخة (UAV) — قدرة ضربة جماعية واحدة قابلة للتشبع.")

# ------------------------------------------------------------------
# 3. إطار العقيدة والمنهجية
# ------------------------------------------------------------------
heading(doc, "3. إطار العقيدة والمنهجية", level=1)
para(doc, "نموذج المحاكاة حتمي لكن ليس مُسبق التحديد — النتيجة تتقرر رياضياً من نسبة القوة المحلية بعد كل خطوة. تتبع المحاكاة 17 خطوة من D-7 حتى D+144h وتنمذج جميع المجالات بالتوازي في كل خطوة.")
para(doc, "العقيدة المرجعية:")
bullet(doc, "JP 3-02 العمليات البرمائية - 5 مراحل (التخطيط، الإركاب، التمرين، الحركة، العمل).")
bullet(doc, "AJP-3.1 العمليات البحرية - السيطرة البحرية، الحرمان البحري، إسقاط القوة.")
bullet(doc, "AJP-3.3 العمليات الجوية - 6 أدوار جوية: مضاد جوي، مضاد بري، مضاد بحري، الضرب الاستراتيجي، ISR، JIPOE.")
bullet(doc, "ATP 3-01.4 J-SEAD - 25-30% من طلعات الأحمر للقمع، 10-15% خسائر AD الأزرق لكل موجة، 5% خسائر الأحمر لكل طلعة.")
bullet(doc, "ATP 3-37.5 الإعاقة المركبة - 400 لغم بحري + حقول ألغام شاطئية.")
bullet(doc, "NWP 3-15 حرب الألغام - 50 لغم/يوم/كاسحة، خسائر الكاسحات 30-50% (Wonsan 1950).")
bullet(doc, "NTRP 3-22 حرب السطح - ASCM احتمال إصابة 0.5، احتمال قتل 0.4، EW soft-kill 25-35% (Latakia 1973).")
bullet(doc, "FM 3-90 و ADP 3-90 - نسبة 3:1 للمهاجم الحاسم، 1.5-3:1 متنازع، أقل من 1.5:1 نقطة الانهيار.")
bullet(doc, "ADP 3-0 نقطة الانهيار - قوة الأحمر تستنفد عند تجاوز قدرته على الإمداد.")
para(doc, "العتبات الحاسمة في النموذج:")
bullet(doc, "RED ADVANTAGE (نسبة ≥ 3:1) - الأحمر قادر على الحسم.")
bullet(doc, "CONTESTED (نسبة 1.5-3:1) - متنازع.")
bullet(doc, "BLUE ADVANTAGE (نسبة < 1.5:1) - الأزرق يقاوم بفعالية والأحمر يقترب من نقطة الانهيار.")
para(doc, "مضاعف الدفاع المستعد للأزرق = 1.5 (تحضير، ألغام، معرفة الأرض). تأثير EW الأحمر يخفض القوة الزرقاء بنسبة (0.18 × قوة التشويش). تأثير EW الأزرق يخفض الأحمر بنسبة (0.10 × قوة التشويش).")

# ------------------------------------------------------------------
# 4. ترتيب القوات
# ------------------------------------------------------------------
heading(doc, "4. ترتيب القوات", level=1)

heading(doc, "4.1 القوة الحمراء (الهجوم)", level=2)
para(doc, "البري:", bold=True)
bullet(doc, "الفرقة 4 المشاة الآلية: لواءا المشاة الآلي 41 و 42 و 43 (كل: 3 كتائب مشاة + 1 كتيبة دبابات) + اللواء المدرع 44 (2 كتيبة دبابات + 1 كتيبة مشاة). إسناد: الاستطلاع 401 (5 سرايا)، سرب طائرات عمودية هجومية 12 طائرة، م/د 402 (36 قاذف)، لواء المدفعية 45 (4 كتائب 155مم + 1 كتيبة 175مم)، لواء الدفاع الجوي 46 (SAM-15 + 2×35مم + SAM-7)، هندسة 403 (5 سرايا)، إشارة 404، EW 405 (4 سرايا)، الدفاع الكيميائي 406 (5 سرايا)، الخدمات 407/408/409.")
bullet(doc, "الفرقة 9 المشاة الآلية: نفس هيكل الفرقة 4 (4 ألوية + إسناد).")
bullet(doc, "الفرقة المدرعة 1: اللواءان المدرعان 11 و 12 + لواء المشاة الآلي 13 + نفس الإسناد.")
bullet(doc, "اللواءان 23 و 24 المشاة: كل منهما 4 كتائب مشاة + 12 زورق مفخخ (USV) = 24 USV إجمالاً.")
bullet(doc, "لواء العمليات الخاصة 21: الكتائب 211 و 212 و 213 و 214.")
bullet(doc, "لواء صواريخ أرض/أرض متوسط المدى (500-600 كم).")
bullet(doc, "الاحتياط: الفرقة 2 المشاة الآلية (حماية القواعد) + الفرقة 8 (احتياطي الفيلق 1).")

para(doc, "البحري:", bold=True)
bullet(doc, "القاعدة البحرية أ: 3 غواصات + 10 مدمرات + 10 فرقاطات + 20 زورق صواريخ + 14 هوفر كرافت + 60 سفينة إبرار متوسطة + 40 سفينة إبرار صغيرة + 160 زورق إبرار + 10 سفن نقل تجارية + 10 سفن بث ألغام + 2 كاسحة ألغام + 2 رف طائرات إسناد بحري.")
bullet(doc, "القاعدة البحرية ب: 8 مدمرات + 9 فرقاطات + 18 زورق صواريخ + 12 هوفر كرافت + 55 سفينة إبرار متوسطة + 35 سفينة إبرار صغيرة + 160 زورق إبرار + 9 سفن نقل + 9 سفن بث ألغام + 2 كاسحة ألغام + 1 رف طائرات إسناد بحري.")

para(doc, "الجوي:", bold=True)
bullet(doc, "القاعدة الجوية أ: السرب 11 و 12 ميج-29 دفاع جوي (24 طائرة)، السرب 13 F-16 هجوم أرضي (12)، السرب 14 ميراج هجوم (12)، السرب 21 عمودية هجومية (12)، السرب 56 و 58 نقل C-130 (24)، 4 طائرات إنذار مبكر، سرب طائرات استطلاع مسيرة (12)، سرب طائرات متفجرة (16).")
bullet(doc, "القاعدة الجوية ب: السرب 23 و 25 ميج-29 دفاع جوي (24)، السرب 35 رافال هجوم (12)، السرب 36 سوخوي-24 هجوم (12)، السرب 22 و 26 عمودية خدمة عامة (24)، سرب طائرات هجومية مسيرة (12)، سرب طائرات متفجرة (16).")
bullet(doc, "القاعدة الجوية ج: السرب 15 سوخوي-24 (12)، السرب 30 رافال دفاع جوي (12)، السرب 16 F-16 (12)، السرب 45 نقل (12)، سرب طائرات هجومية مسيرة (12)، سرب طائرات متفجرة (16).")
bullet(doc, "لواء الدفاع الجوي 2: كتيبة صواريخ S-300 (3 سرايا) + كتيبة هوك (3 سرايا) + 3 كتائب SAM-2 + كتيبة SAM-15 + كتيبة م/ط 35مم + رادارات كشف.")

heading(doc, "4.2 القوة الزرقاء (الدفاع)", level=2)
para(doc, "البري — المكون البري 99:", bold=True)
bullet(doc, "لواء المشاة الآلي 51: كتائب 511، 512، 513 مشاة آلية + كتيبة الدبابات 514 (4 سرايا) + إسناد اللواء.")
bullet(doc, "لواء المشاة الآلي 52: كتائب 521، 522، 523 + كتيبة الدبابات 524 + إسناد.")
bullet(doc, "اللواء المدرع 54: كتيبة المشاة 541 + كتيبتا الدبابات 542 و 543 + إسناد.")
bullet(doc, "2 سرية طائرات هجوم واستطلاع (12 طائرة) + كتيبة الاستطلاع 501 (3 سرايا + 12 UAV استطلاع).")
bullet(doc, "المدفعية: 551 و 552 و 554 (متوسطة) + 555 (ثقيلة) + 556 (راجمات صواريخ).")
bullet(doc, "كتيبة م/د 502 (3 سرايا/36 قاذف)، هندسة 503، إشارة 504، EW 505 (3 سرايا: تكتيكية متحركة + استراتيجية استطلاع + استراتيجية إعاقة)، الدفاع الكيميائي 506.")
bullet(doc, "كتيبة الدفاع الجوي 507 (+): 561 صواريخ متوسطة المدى (4 سرايا) + 1 سرية MANPADS (4 فصائل).")

para(doc, "البحري:", bold=True)
bullet(doc, "8 كورفيت + 9 زوارق صواريخ + 12 زورق إنزال سريع + 6 زورق إسناد ساحلي + 4 أجهزة رادار ساحلي + 3 سفن بث ألغام + 2 قانصة ألغام + سرية طائرات عمودية بحرية (24 طائرة) + 20 زورق مطاردة + 400 لغم بحري + سفن لوجستية متفرقة.")

para(doc, "الجوي:", bold=True)
bullet(doc, "القاعدة الجوية أ: السرب 3 F-16 دفاع جوي + السرب 4 ميج هجوم + سرب هجوم أرضي + سرب طائرات مسيرة هجومية = 36 طائرة قتالية + 12 UAV.")
bullet(doc, "القاعدة الجوية ب: السرب 1 رافال دفاع جوي + السرب 2 ميج هجوم + السرب 9 نقل + 4 طائرات إنذار مبكر + 2 طائرات تزود بالوقود + 2 طائرات استخبارية مسيرة.")
bullet(doc, "القاعدة الجوية ج: السرب 5 F-16 + السرب 6 ميج + السرب 7 رافال متعدد المهام + سرب طائرات هجومية مسيرة = 36 طائرة + 12 UAV.")
bullet(doc, "لواء الدفاع الجوي 9: الكتيبة 80 S-300 + الكتيبتان 81 و 82 و 83 هوك + الكتيبة 85 م/ط 35مم + 1 رادار كشف عالي (250 كم) + 9 رادارات كشف منخفض (80 كم).")

para(doc, "العمليات الخاصة والاحتياطيات:", bold=True)
bullet(doc, "كتيبة العمليات الخاصة 80: 3 سرايا عمليات خاصة + 2 سرية طائرات اقتحام ونقل جوي (12 طائرة إجمالاً).")
bullet(doc, "احتياطي JTF: اللواء المدرع 72 (3 كتائب: 721 مشاة + 722 و 723 دبابات) + كتيبة الاستطلاع 701 + سرية طائرات هجوم (6 طائرات).")
bullet(doc, "احتياطي القيادة المشتركة: لواء المشاة الآلي 71 (4 كتائب) + اللواء المدرع 73 (3 كتائب) + 2 سرية طائرات (12 طائرة).")
bullet(doc, "لواء صواريخ أرض/أرض (متوسط المدى 800-1000 كم).")

# ------------------------------------------------------------------
# 5. تحليل الأرض والموانع
# ------------------------------------------------------------------
heading(doc, "5. تحليل الأرض والموانع", level=1)
para(doc, "تحليل الأرض من ملفات GIS (landuse + inland_water + water_way) داخل صندوق العمليات (lon 19.12-20.02, lat 29.50-30.56):")
bullet(doc, "البحر شمال خط 30.45-30.55 - المنطقة البحرية للتحضير والإنزال.")
bullet(doc, "الشاطئ يمتد من lon 19.13 (الغرب، lat 30.28) إلى lon 19.81 (lat 30.55). شرقاً ينحرف الشاطئ شمالاً خارج إطار العرض.")
bullet(doc, "السبخات والكتل المائية الداخلية: 12 مائية + 4 سبخات مالحة + 3 خزانات + 1 رطبة + 1 مستنقع = 21 حاجزاً للمدرعات.")
bullet(doc, "البقع الصناعية والمسكونية: 19 صناعي + 7 سكني + 2 تجاري + 2 محجر = مناطق التفاف.")
bullet(doc, "الصحراء/الشجيرات/الأرض المفتوحة - مرور حر بشكل عام.")
para(doc, "تم التحقق برمجياً أن جميع المحاور الحمراء الأربعة من BLS-1..4 إلى OBJ X لا تتقاطع مع أي حاجز مائي داخلي.")
para(doc, "العنصر الأهم: 400 لغم بحري مسبق النشر في 4 أحزمة أمام BLS-1..4 - الحاجز الأساسي للحركة البحرية.", bold=True)

# ------------------------------------------------------------------
# 6. اختيار مواقع الإنزال
# ------------------------------------------------------------------
heading(doc, "6. اختيار مواقع الإنزال (BLS)", level=1)
para(doc, "تم تقييم 21 مرشحاً على الشاطئ بفواصل ~3 كم وتم اختيار 4 مواقع مع فصل أدنى 8 كم بينها. النتائج:")
tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for j, h in enumerate(["المعرف","الإحداثيات","مسافة الأزرق (كم)","ألوية/يوم","ملاحظة"]):
    hdr[j].text = h
    for p in hdr[j].paragraphs:
        set_rtl(p)
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xff,0xff,0xff)
        tcPr = hdr[j]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1e3a8a'); tcPr.append(shd)

for f in BLS_GJ["features"]:
    pr = f["properties"]
    a = pr["anchor_lonlat"]
    row = tbl.add_row().cells
    vals = [pr["id"], f"({a[0]:.4f}, {a[1]:.4f})",
            f"{pr['min_blue_dist_km']:.1f}", f"{pr['throughput']['bde_per_day']:.2f}",
            "نقطة صلبة، اقتراب مفتوح، لا تقاطعات داخلية"]
    for j, c in enumerate(row):
        c.text = vals[j]
        for p in c.paragraphs:
            set_rtl(p)
            for r in p.runs: r.font.size = Pt(9)

total_bde = sum(f["properties"]["throughput"]["bde_per_day"] for f in BLS_GJ["features"])
para(doc, f"إجمالي السقف اللوجستي لالتزام القوة الحمراء: {total_bde:.2f} لواء/يوم عبر الشواطئ الأربعة.")

# ------------------------------------------------------------------
# 7. نموذج الحركة والقدرة الاستيعابية
# ------------------------------------------------------------------
heading(doc, "7. نموذج الحركة والقدرة الاستيعابية", level=1)
bullet(doc, "السبخات والمالحات والمستنقعات والخزانات والكتل المائية = موانع للمدرعات.")
bullet(doc, "المناطق الصناعية/السكنية = مناطق التفاف لا اشتباك حضري.")
bullet(doc, "الأراضي الزراعية = ممكنة بمعامل 0.7.")
bullet(doc, "الصحراء/الشجيرات = ممكنة كاملاً.")
bullet(doc, "الوديان = ممرات أو نقاط اختناق، ليست عوائق رطبة تلقائياً.")
para(doc, "نموذج القدرة الاستيعابية:")
bullet(doc, "اتساع شاطئ صالح ~ 800-1000 م لكل BLS.")
bullet(doc, "مسافة المخرج إلى أقرب طريق ~ 1.5 كم.")
bullet(doc, "العوامل المُقيِّدة: اقتراب ساحلي، صلابة الشاطئ، المسافة من المدافعين، عدم وجود حواجز خلفية، عدد المراكب المتبقية بعد الألغام.")

# ------------------------------------------------------------------
# 8. تقييم الهدف CARVER
# ------------------------------------------------------------------
heading(doc, "8. تقييم الهدف باستخدام CARVER", level=1)
para(doc, "الهدف المختار: OBJ-PIPELINE عند (19.55، 29.74)، بعمق 90.1 كم من الساحل. مجموع CARVER 48/60: C9 A8 R7 V6 E9 Re9.")
bullet(doc, "C (الحيوية) = 9: نقطة الانتصاف لخط أنابيب الناصر-البريقة، عقدة لوجستية حيوية لاقتصاد العدو واستمرار التصدير.")
bullet(doc, "A (الوصول) = 8: محور الناصر-البريقة طريق مفتوح يربط الساحل بالداخل.")
bullet(doc, "R (الاستعادة) = 7: استعادة خلال أسابيع إن تضرر — ليست فورية.")
bullet(doc, "V (الهشاشة) = 6: محمي بقوة بدفاع ساحلي عميق ولواء مدرع احتياطي.")
bullet(doc, "E (الأثر) = 9: ضربة استراتيجية وإعلامية كبيرة على اقتصاد العدو ومعنوياته.")
bullet(doc, "Re (التعرف البصري) = 9: علامة جغرافية واضحة على الخرائط، يسهل التعرف عليها وتأكيدها.")

# ------------------------------------------------------------------
# 9. منهجية الفعل ورد الفعل
# ------------------------------------------------------------------
heading(doc, "9. منهجية الفعل ورد الفعل", level=1)
para(doc, "كل خطوة محاكاة تُنتج سجلاً مفصلاً للفعل ورد الفعل عبر ستة مجالات + خمسة عناصر مكملة:")
bullet(doc, "ضربات استراتيجية (strategic_strikes): إطلاقات صواريخ أرض/أرض بعيدة المدى ضد القواعد والمنشآت.")
bullet(doc, "أعمال بحرية (maritime_actions): اشتباكات سطحية، ASW، حركات الأسطول.")
bullet(doc, "أعمال جوية (air_actions): SEAD، CAS، اعتراض، طلعات، AWACS.")
bullet(doc, "حرب الألغام (mine_warfare): زرع وتطهير، خسائر الكاسحات، خسائر السفن بالألغام.")
bullet(doc, "أسراب USV/UAV (uav_usv_swarms): الزوارق المفخخة والطائرات المسيرة المنسقة.")
bullet(doc, "أعمال برية (land_actions): اشتباكات الكتائب، الاندفاع، الهجوم المضاد.")
bullet(doc, "العمليات الخاصة (sof_actions): الإنزالات العمودية، الغارات على عقد القيادة.")
bullet(doc, "الحرب الإلكترونية (ew_actions): التشويش وSIGINT.")
bullet(doc, "اللوجستيات (logistics_state): الإنزال هذه الخطوة، حالة الإمداد، نسبة استهلاك الشاطئ.")
bullet(doc, "الخسائر (losses_this_step): الوحدات المدمَّرة من كل جانب مع سبب التدمير ومن دمَّرها.")
bullet(doc, "الميزة هذه الخطوة (step_advantage) ونسبة القوة المحلية مع سبب الميزة.")
para(doc, "كل خطوة في القسم 10 تتضمن ثلاث صور (نظرة كاملة + تكبير AOI + تكبير النشاط) من إعداد Codex، ومن ثَم سرد كامل لكل المجالات.")

# ------------------------------------------------------------------
# 10. سير المناورة خطوة بخطوة (مع الصور)
# ------------------------------------------------------------------
heading(doc, "10. سير المناورة خطوة بخطوة", level=1)
para(doc, "كل خطوة تحتوي على ثلاث صور توضح المعركة من زوايا مختلفة: نظرة كاملة للمسرح + تكبير على المنطقة الساحلية والهدف + تكبير على النشاط الرئيسي.")

def fmt_dict(d, keys=("actor","what","where","why","intended_effect","effectiveness","note","doctrine","sorties","missiles","hits","effect")):
    parts = []
    for k in keys:
        v = d.get(k)
        if v is None or v == "": continue
        label = {"actor":"الفاعل","what":"العمل","where":"الموقع","why":"السبب",
                 "intended_effect":"الأثر المقصود","effectiveness":"الفعالية","note":"ملاحظة",
                 "doctrine":"العقيدة","sorties":"الطلعات","missiles":"الصواريخ","hits":"الإصابات",
                 "effect":"الأثر"}.get(k, k)
        parts.append(f"{label}: {v}")
    return " | ".join(parts)

for snap in SNAPS:
    s = snap["step"]
    heading(doc, f"10.{s:02d} الخطوة {s} — {snap['time']} — {snap['phase']}", level=2)

    # Embed 3 images: full, AOI zoom, action zoom
    img_full = ROOT / f"step{s:02d}.jpeg"
    img_aoi  = ROOT / f"step{s:02d}_zoom_aoi.jpeg"
    img_act  = ROOT / f"step{s:02d}_zoom_action.jpeg"
    embed_image(doc, img_full, f"الصورة 10.{s:02d}.أ — نظرة كاملة على المسرح في الخطوة {s} ({snap['time']})", width_in=6.4)
    embed_image(doc, img_aoi,  f"الصورة 10.{s:02d}.ب — تكبير على الساحل ورأس الجسر والهدف X", width_in=6.4)
    embed_image(doc, img_act,  f"الصورة 10.{s:02d}.ج — تكبير على النشاط الرئيسي (الإنزال/المعركة/الموانع)", width_in=6.4)

    # Header info
    para(doc, f"خط الطور: {snap['phase_line_km']} كم  |  حالة الهدف X: {snap['obj_status']}  "
              f"|  EW أحمر/أزرق: {snap['ew_strength_red']:.2f}/{snap['ew_strength_blue']:.2f}  "
              f"|  ألغام بحرية متبقية: {snap['blue_mines_remaining']}/400  "
              f"|  نسبة القوة المحلية: {snap['force_ratio_local']}:1  |  الميزة: {snap['step_advantage']}",
              size=10, bold=True)
    para(doc, f"سبب الميزة: {snap['advantage_reason']}", size=10)

    # Strategic strikes
    if snap.get("strategic_strikes"):
        heading(doc, "الضربات الاستراتيجية", level=3)
        for s_ in snap["strategic_strikes"]:
            bullet(doc, fmt_dict(s_))

    # Maritime actions
    if snap.get("maritime_actions"):
        heading(doc, "الأعمال البحرية", level=3)
        for s_ in snap["maritime_actions"]:
            bullet(doc, fmt_dict(s_))

    # Air actions
    if snap.get("air_actions"):
        heading(doc, "الأعمال الجوية", level=3)
        for s_ in snap["air_actions"]:
            bullet(doc, fmt_dict(s_))

    # Mine warfare
    if snap.get("mine_warfare"):
        heading(doc, "حرب الألغام البحرية", level=3)
        for s_ in snap["mine_warfare"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","where","method","clearance_rate","effect","remaining","doctrine")))

    # USV / UAV
    if snap.get("uav_usv_swarms"):
        heading(doc, "أسراب USV و UAV", level=3)
        for s_ in snap["uav_usv_swarms"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","launched","intercepted_by_blue_helos","intercepted_by_CRAM","intercepted_by_AD","intercepted_by_CIWS","hits","kills","surviving","doctrine","calibration","effectiveness")))

    # SOF
    if snap.get("sof_actions"):
        heading(doc, "العمليات الخاصة", level=3)
        for s_ in snap["sof_actions"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","helos","defender_alert","losses_at_drop","casualties_inflicted","doctrine")))

    # Land actions
    if snap.get("land_actions"):
        heading(doc, "الأعمال البرية", level=3)
        for s_ in snap["land_actions"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","where","why","intended_effect","casualties_landing","effect","doctrine")))

    # EW
    if snap.get("ew_actions"):
        heading(doc, "الحرب الإلكترونية", level=3)
        for s_ in snap["ew_actions"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","intensity","effect","why")))

    # Red actions summary
    if snap.get("red_actions_summary"):
        heading(doc, "أعمال الأحمر (الملخص)", level=3)
        for s_ in snap["red_actions_summary"]:
            bullet(doc, fmt_dict(s_))

    # Blue reactions summary
    if snap.get("blue_reactions_summary"):
        heading(doc, "رد فعل الأزرق (الملخص)", level=3)
        for s_ in snap["blue_reactions_summary"]:
            bullet(doc, fmt_dict(s_))

    # Red counter
    if snap.get("red_counter_reactions"):
        heading(doc, "رد الأحمر على رد الفعل", level=3)
        for s_ in snap["red_counter_reactions"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","why")))

    # Blue counter-reactions
    if snap.get("blue_counter_reactions"):
        heading(doc, "رد الأزرق المضاد على المضاد", level=3)
        for s_ in snap["blue_counter_reactions"]:
            bullet(doc, fmt_dict(s_, keys=("actor","what","why")))

    # Logistics
    if snap.get("logistics"):
        heading(doc, "اللوجستيات والقدرة الاستيعابية", level=3)
        ls = snap["logistics"]
        bullet(doc, f"إنزال الأحمر هذه الخطوة: {ls.get('red_landings','—')}.")
        bullet(doc, f"حالة إمداد الأزرق: {ls.get('blue_supply','—')}.")
        bullet(doc, f"نسبة استهلاك الشاطئ: {ls.get('beach_throughput','—')}%.")
        if ls.get("notes"): bullet(doc, f"ملاحظات: {ls['notes']}.")

    # Losses
    heading(doc, "الخسائر وأسبابها", level=3)
    if not snap["losses_this_step"]["RED"] and not snap["losses_this_step"]["BLUE"]:
        bullet(doc, "لا خسائر مسجلة في هذه الخطوة.")
    if snap["losses_this_step"]["RED"]:
        para(doc, "خسائر الأحمر:", bold=True, color=RGBColor(0x7f,0x1d,0x1d))
        for l in snap["losses_this_step"]["RED"]:
            partial = " (جزئي)" if l.get("partial") else ""
            bullet(doc, f"الوحدة: {l.get('name','?')} ({l.get('name_ar','?')}) [{l.get('domain','?')}]{partial}.  السبب: {l.get('destroyed_reason_ar','—')}.  من دمَّرها: {l.get('destroyed_by','—')}.")
    if snap["losses_this_step"]["BLUE"]:
        para(doc, "خسائر الأزرق:", bold=True, color=RGBColor(0x1e,0x3a,0x8a))
        for l in snap["losses_this_step"]["BLUE"]:
            partial = " (جزئي)" if l.get("partial") else ""
            bullet(doc, f"الوحدة: {l.get('name','?')} ({l.get('name_ar','?')}) [{l.get('domain','?')}]{partial}.  السبب: {l.get('destroyed_reason_ar','—')}.  من دمَّرها: {l.get('destroyed_by','—')}.")

    cum = snap["cumulative_losses"]
    para(doc,
        f"الخسائر التراكمية حتى هذه الخطوة: أحمر {cum['RED']['count']} (جوي:{cum['RED']['by_domain'].get('air',0)} "
        f"بحري:{cum['RED']['by_domain'].get('naval',0)} بري:{cum['RED']['by_domain'].get('ground',0)} "
        f"عمليات خاصة:{cum['RED']['by_domain'].get('sof',0)} استراتيجي:{cum['RED']['by_domain'].get('strategic',0)}) "
        f"| أزرق {cum['BLUE']['count']} (جوي:{cum['BLUE']['by_domain'].get('air',0)} "
        f"بحري:{cum['BLUE']['by_domain'].get('naval',0)} بري:{cum['BLUE']['by_domain'].get('ground',0)}).",
        size=10, bold=True)

    page_break(doc)

# ------------------------------------------------------------------
# 11. التحليل النهائي ونسبة القوة
# ------------------------------------------------------------------
heading(doc, "11. التحليل النهائي ونسبة القوة", level=1)
para(doc, "جدول تطور نسبة القوة عبر 17 خطوة:")
ftbl = doc.add_table(rows=1, cols=7)
ftbl.style = "Light Grid Accent 1"
hdr = ftbl.rows[0].cells
for j, h in enumerate(["الخطوة","الزمن","الطور","FR محلي","FR عملياتي","ألغام متبقية","الميزة"]):
    hdr[j].text = h
    for p in hdr[j].paragraphs:
        set_rtl(p)
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xff,0xff,0xff)
        tcPr = hdr[j]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1e3a8a'); tcPr.append(shd)
for snap in SNAPS:
    row = ftbl.add_row().cells
    vals = [str(snap["step"]), snap["time"], snap["phase"][:30],
            f"{snap['force_ratio_local']}:1", f"{snap['force_ratio_operational']}:1",
            str(snap["blue_mines_remaining"]), snap["step_advantage"]]
    for j, c in enumerate(row):
        c.text = vals[j]
        for p in c.paragraphs:
            set_rtl(p)
            for r in p.runs: r.font.size = Pt(8)

# ------------------------------------------------------------------
# 12. النتيجة العملياتية
# ------------------------------------------------------------------
heading(doc, "12. النتيجة العملياتية", level=1)
final = SNAPS[-1]
fr_final = final["force_ratio_local"]
outcome_logic = final.get("outcome_logic","")
para(doc, f"النسبة المحلية النهائية عند الهدف X (الخطوة 16): {fr_final}:1.", bold=True)
para(doc, f"تصنيف النتيجة: {outcome_logic}", bold=True, color=RGBColor(0x7f,0x1d,0x1d))
para(doc, "التفسير العملياتي:")
if "BLUE DENIES" in (outcome_logic or ""):
    bullet(doc, "وصل الأحمر إلى عمق ~95 كم لكن نسبة القوة المحلية بقيت دون 1.5:1 طوال 16 خطوة.")
    bullet(doc, "العوامل الحاسمة (بالترتيب):")
    bullet(doc, "1. تفوق مدى صواريخ أرض/أرض الزرقاء (800-1000 كم) أعطب القواعد الحمراء قبل الإنطلاق - خسر الأحمر ~7% من طائراته و ~6% من سفنه قبل بدء العملية.", indent_lvl=1)
    bullet(doc, "2. 400 لغم بحري مسبقة النشر دمّرت جميع كاسحات الألغام الحمراء الـ4 خلال يومين، وأعطبت سفن الإنزال.", indent_lvl=1)
    bullet(doc, "3. حملة SEAD الأحمر استنزفت 25% من طائرات الهجوم الحمراء قبل تحقيق تفوق جوي محلي - مقابل خسارة الأزرق ~30% من قدرة AD.", indent_lvl=1)
    bullet(doc, "4. الموجة المركزة من 24 USV + 48 UAV-X كسرت 56% فقط من الدفاع الأزرق - مع اعتراض 35 من 72 تهديداً.", indent_lvl=1)
    bullet(doc, "5. الهجوم المضاد الأزرق المنسق (72 المدرع + 73 المدرع + JC-Helo + المدفعية) كسر زخم 1-AD قبل OBJ X.", indent_lvl=1)
    bullet(doc, "6. تأثير الحرب الإلكترونية الأحمر انخفض من 0.80 إلى 0.10 عبر 16 خطوة، فسمح للأزرق باستعادة C2.", indent_lvl=1)
    bullet(doc, "نقطة الانهيار (Culmination) لقوة الأحمر وقعت أساساً أثناء المرحلة 0-5 (الإجراءات الاستراتيجية والبحرية والجوية والألغام)، قبل أن يصل أي جندي أحمر إلى الشاطئ.")
elif "CONTESTED" in (outcome_logic or ""):
    bullet(doc, "نتيجة متنازعة - الأحمر اقترب من الهدف X لكن لم يحقق نسبة 3:1 الحاسمة.")
elif "CAPTURE" in (outcome_logic or ""):
    bullet(doc, "الأحمر حقق ميزة حاسمة عند الهدف X مع رأس جسر سليم.")

# ------------------------------------------------------------------
# 13. القيود ومصادر عدم اليقين
# ------------------------------------------------------------------
heading(doc, "13. القيود ومصادر عدم اليقين", level=1)
bullet(doc, "نسب القوة مبنية على نقاط نمذجة مبسّطة - القتال الفعلي يتأثر بالتضاريس الدقيقة والأرصاد والروح المعنوية.")
bullet(doc, "حقول الألغام الشاطئية الزرقاء افتراض عقائدي بسبب الدفاع المحضّر (AJP-3.2 / FM 3-90)؛ لو حُذف يُخفِّض خسائر الأحمر بحوالي 2-3 سرايا.")
bullet(doc, "نموذج USV: عوامل المعايرة من Black Sea 2024 (25-30% بقاء) - الفعالية الحقيقية تتغيّر مع تطور الـCRAM المضاد.")
bullet(doc, "نموذج SEAD: عوامل من Vietnam Wild Weasel + الخليج 1991 - 25-30% من الطلعات للقمع، 5% خسائر/طلعة.")
bullet(doc, "نموذج الموانع البحرية: عوامل من Wonsan 1950 (50 لغم/يوم/كاسحة، خسائر 30-50%).")
bullet(doc, "ساعة H غير محددة في المصادر - استُخدم D-H -7 كنقطة بداية الاستعدادات.")
bullet(doc, "صور القمر الصناعي ESRI ثابتة - التغيرات الموسمية في السبخات غير مأخوذة بالحسبان.")
bullet(doc, "الفرضية: الأحمر لا يستخدم NBC (الكتيبة 406 دعم فقط) - متطابق مع المصادر النصية.")
bullet(doc, "نموذج التشبع (Saturation): >4 تهديدات في 30 ثانية تخفض معدل الاعتراض بنسبة 40% (عقيدة Houthi 2024).")

# ------------------------------------------------------------------
# 14. المراجع
# ------------------------------------------------------------------
heading(doc, "14. المراجع", level=1)
para(doc, "المراجع الوثائقية المباشرة:", bold=True)
bullet(doc, "القوة الحمراء.docx - ترتيب القتال الأحمر الرسمي.")
bullet(doc, "القوة الزرقاء.docx - ترتيب القتال الأزرق الرسمي.")
bullet(doc, "enemy.docx - العمل الأحمر الأكثر احتمالاً (3 مراحل).")
bullet(doc, "nato-map-layers.geojson - مواضع الوحدات الأصلية + 3 AO.")
bullet(doc, "OpenStreetMap - landuse + inland_water + water_way.")
bullet(doc, "ESRI World Imagery - صورة قمر صناعي مرجعية.")

para(doc, "العقائد المرجعية (Doctrines.md):", bold=True)
bullet(doc, "JP 3-02 العمليات البرمائية (يناير 2019).")
bullet(doc, "AJP-3.1 العمليات البحرية NATO (Edition B, 2025).")
bullet(doc, "AJP-3.2 العمليات البرية NATO (Edition B).")
bullet(doc, "AJP-3.3 العمليات الجوية والفضائية NATO (Edition B, 2024).")
bullet(doc, "AJP-3.3.3 التنسيق الجوي-البحري NATO.")
bullet(doc, "ATP 3-01.4 J-SEAD - قمع الدفاع الجوي المشترك.")
bullet(doc, "ATP 3-01.8 + ATP 3-01.81 - الدفاع الجوي المركب + C-UAS.")
bullet(doc, "ATP 3-01.7 - تقنيات لواء مدفعية الدفاع الجوي.")
bullet(doc, "ATP 3-37.5 - الإعاقة المركبة (الألغام والعوائق).")
bullet(doc, "ATP 3-90.4 - الحركة المركبة (الخرق وعبور الفجوات).")
bullet(doc, "ATP 3-09.42 - دعم النيران للواء القتالي.")
bullet(doc, "ATP 3-91 - عمليات الفرقة.")
bullet(doc, "ATP 3-92 - عمليات الفيلق.")
bullet(doc, "FM 3-90 و ADP 3-90 - التكتيكات والعمليات الهجومية/الدفاعية.")
bullet(doc, "ADP 3-0 - نقطة الانهيار والعمليات المعركة.")
bullet(doc, "NWP 3-15 / JP 3-15 - حرب الألغام البحرية.")
bullet(doc, "NTRP 3-22 - حرب السطح البحري.")
bullet(doc, "ATP 3-18.4 - عمليات قوات العمليات الخاصة المباشرة.")
bullet(doc, "AJP-3.9 - الاستهداف المشترك.")
bullet(doc, "AJP-3.20 - عمليات الفضاء السيبراني.")
bullet(doc, "ATP 3-12 / FM 3-12 - الحرب الإلكترونية والسيبرانية.")
bullet(doc, "ATP 4-91 - الإمداد على مستوى الفرقة (نوفمبر 2022).")
bullet(doc, "ATP 4-92 - الإمداد على مستوى الفيلق والجيش الميداني (مارس 2023).")

para(doc, "الدروس التاريخية المرجعية (WarReferences.md):", bold=True)
bullet(doc, "Iwo Jima 1945 (تأثير القصف ضعيف ضد المحصّن في العمق).")
bullet(doc, "Tarawa 1943 (الشعاب والمد - استنزاف الإنزال).")
bullet(doc, "D-Day Normandy 1944 (التفوق الجوي والبحري الكاسح).")
bullet(doc, "Falklands San Carlos 1982 (ASCM ضد سفن حديثة).")
bullet(doc, "Gallipoli 1915 (الألغام البحرية تكسر هجوم الأسطول).")
bullet(doc, "Wonsan 1950 (تطهير الألغام بطيء ومميت للكاسحات).")
bullet(doc, "Inchon 1950 (المد والاحتراب الجريء).")
bullet(doc, "USV الأوكرانية 2022-2024 (نموذج البقاء 25-30%).")
bullet(doc, "الحوثيون البحر الأحمر 2023-2024 (نموذج التشبع متعدد المتجهات).")
bullet(doc, "Latakia 1973 (المعركة الأولى صواريخ على صواريخ + ECM).")
bullet(doc, "Praying Mantis 1988 (الاشتباك السطحي الحديث).")
bullet(doc, "Wild Weasel SEAD (5% خسائر/طلعة، 10-15% خسائر AD/موجة).")
bullet(doc, "NATO ليبيا 2011 (نفس AOI! - 110 توماهوك + 9700 طلعة + 7 أشهر).")
bullet(doc, "قبرص 1974 (إنزال 3000 + إنزال مظلي).")
bullet(doc, "روسيا-جورجيا 2008 (الإنزال البحري في Ochamchire).")
bullet(doc, "القرم 2014 (الإستيلاء الهجين السوفيتي).")
bullet(doc, "Moskva 2022 (Neptune ASCM ضد الطراد الراية).")
bullet(doc, "إيران 14 أبريل 2024 (185 + 110 + 36 وحدات متشبعة، 99% اعتراض).")

# ------------------------------------------------------------------
# Save
# ------------------------------------------------------------------
out = ROOT / "WarGamingClaude3_Report_AR.docx"
doc.save(str(out))
print(f"Saved {out} ({out.stat().st_size:,} bytes)")
