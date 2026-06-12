"""
Build a polished Arabic operation-summary DOCX for Wargame1 only.

The report is intentionally written as a finished operation/wargame report. It
does not discuss previous drafts or correction history.
"""

from __future__ import annotations

from pathlib import Path
import json

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/hextechkraken/Desktop/TestingAI")
DIR = ROOT / "Wargame1"
OUT = DIR / "Wargame1_Report_AR.docx"

BLUE = RGBColor(31, 56, 100)
MID_BLUE = RGBColor(46, 83, 149)
DARK = RGBColor(25, 31, 40)
MUTED = RGBColor(92, 100, 112)
RED = RGBColor(150, 30, 30)
GOLD = RGBColor(122, 90, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        rpr = run._r.get_or_add_rPr()
        rtl = rpr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            rpr.append(rtl)
        rtl.set(qn("w:val"), "1")


def ar_para(doc: Document, text: str = "", *, size=11, bold=False, color=DARK,
            before=0, after=6, style=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = align
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    set_rtl(p)
    return p


def heading(doc: Document, text: str, level=1):
    size = {1: 16, 2: 13, 3: 12}[level]
    color = MID_BLUE if level in (1, 2) else BLUE
    before = {1: 16, 2: 12, 3: 8}[level]
    after = {1: 8, 2: 6, 3: 4}[level]
    p = ar_para(doc, text, size=size, bold=True, color=color, before=before, after=after)
    p.style = f"Heading {level}"
    set_rtl(p)
    return p


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def table(doc: Document, rows, widths=None, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    if widths is None:
        widths = [9360 // len(rows[0])] * len(rows[0])
    for i, row in enumerate(rows):
        tr = t.rows[i]
        for j, val in enumerate(row):
            cell = tr.cells[j]
            set_cell_width(cell, widths[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0 and header:
                set_cell_shading(cell, "E8EEF5")
            for p in cell.paragraphs:
                p.text = ""
                run = p.add_run(str(val))
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
                run.font.size = Pt(9.5)
                run.font.bold = i == 0 and header
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(2)
                set_rtl(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    cap = ar_para(doc, caption, size=9, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def load_steps():
    steps = []
    for idx in range(12):
        with (DIR / f"step{idx:02d}.geojson").open(encoding="utf-8") as f:
            steps.append(json.load(f)["metadata"])
    return steps


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        st.font.bold = True
        st.font.color.rgb = MID_BLUE if style_name != "Heading 3" else BLUE


def build() -> None:
    steps = load_steps()
    doc = Document()
    configure_doc(doc)

    ar_para(doc, "تقرير محاكاة لعبة الحرب", size=28, bold=True, color=BLUE,
            before=100, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    ar_para(doc, "عملية إبرار برمائي - ساحل البريقة / إجدابيا (ليبيا)",
            size=16, bold=True, color=MID_BLUE, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    ar_para(doc, "لعبة الحرب رقم 1", size=14, bold=True, color=MUTED,
            after=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    table(doc, [
        ["البند", "القيمة"],
        ["النموذج", "محاكاة تشغيلية حتمية بمنطق نسبة القوة والتضاريس والتوقيت"],
        ["السيناريو", "هجوم برمائي أحمر وفق COA #1 على ساحل البريقة / إجدابيا"],
        ["الهدف النهائي", "OBJ NASSER-95 على خط أنابيب ناصر-البريقة ضمن نطاق 80-100 كم"],
        ["زمن التشغيل", "من ساعة الصفر إلى س+144 (ستة أيام تشغيلية)"],
        ["حالة النتيجة", "استيلاء الفريق الأحمر على الهدف بعد استنزاف احتياط الفريق الأزرق"],
    ], widths=[1900, 7460])
    ar_para(doc, "تاريخ التحرير: مايو 2026", size=10, color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    heading(doc, "1. المقدمة العملياتية", 1)
    ar_para(doc,
            "تعرض هذه الوثيقة خلاصة تشغيل لعبة حرب لعملية إبرار برمائي على ساحل البريقة / إجدابيا. "
            "الفريق الأزرق يحتفظ بمنطقة عمليات ساحلية ذات ثلاث مناطق فرعية، وينتشر بقوة دفاعية من 39 وحدة "
            "تشمل قيادة فرقة، ثلاثة ألوية، تسع كتائب، وستاً وعشرين سرية. الفريق الأحمر ينفذ خطة تعرضية من "
            "ثلاث مراحل: استطلاع وتأمين رأس الشاطئ، إدخال فرقة المشاة الآلية الرابعة ثم فرقة المشاة الآلية "
            "التاسعة، ثم استثمار الفرقة المدرعة الأولى باتجاه الهدف ناصر.")
    ar_para(doc,
            "الهدف من التشغيل هو تقييم قدرة الفريق الأحمر على تحويل الإبرار الساحلي إلى اختراق بري عميق، "
            "ومدى قدرة الاحتياط الأزرق على تعطيل هذا الاستثمار قبل بلوغ عمق 80-100 كم.")

    heading(doc, "2. الإطار العقائدي ونموذج التحكيم", 1)
    ar_para(doc,
            "يعتمد التشغيل على إطار تحكيم حتمي يوازن بين نسبة القوة المحلية، صلاحية ممرات الإبرار، تأثير "
            "الحرب الإلكترونية، قدرة الخروج من الشاطئ، وتوقيت الاحتياط. لا تمثل الأرقام حساب استنزاف آلياً "
            "مستقلاً، بل قراراً تحكيمياً منظماً لكل مرحلة.")
    table(doc, [
        ["عامل التحكيم", "كيف استُخدم في التشغيل"],
        ["نسبة القوة", "اعتُبرت 3:1 تقريباً عتبة ملائمة لاختراق دفاع مُحضّر، مع نتائج متنازعة بين 1.5:1 و3:1."],
        ["التضاريس", "قيّدت السبخات والبحيرات الساحلية الخروج من بعض نقاط الإبرار، بينما وفرت الصحراء ممرات مناورة أعمق."],
        ["الحرب الإلكترونية", "خفضت تنسيق الفريق الأزرق في المرحلتين الأولى والثانية، ثم تراجع أثرها في العمق."],
        ["اللوجستيات", "تحول الشاطئ من نقطة إنزال إلى منطقة صيانة قتال، مع بقاء القدرة محدودة في BLS-4."],
        ["الاحتياط", "حُسب توقيت التزام الاحتياط الأزرق كعامل حاسم في إبطاء أو كسر الاستثمار الأحمر."],
    ], widths=[2300, 7060])

    heading(doc, "3. نظام المعركة", 1)
    heading(doc, "3.1 الفريق الأحمر", 2)
    table(doc, [
        ["التشكيل", "الدور في التشغيل"],
        ["كتيبة الاستطلاع 401", "تأمين مواضع الإبرار الأولية وتحديد ممرات الخروج."],
        ["فرقة المشاة الآلية 4", "موجة الاقتحام الرئيسية عبر BLS-1 إلى BLS-4."],
        ["فرقة المشاة الآلية 9", "قوة المتابعة لتوسيع رأس الشاطئ والوصول إلى 40-50 كم."],
        ["الفرقة المدرعة 1", "قوة الاستثمار باتجاه OBJ NASSER-95."],
        ["لواء المدفعية 45 / كتيبة الحرب الإلكترونية 405 / 24 زورقاً مسيراً متفجراً", "إسناد ناري، إعاقة قيادة وسيطرة، وتعطيل مواضع الساحل في بداية العملية."],
    ], widths=[3500, 5860])
    heading(doc, "3.2 الفريق الأزرق", 2)
    ar_para(doc,
            "يتكون الدفاع الأزرق من قيادة فرقة، ثلاثة ألوية، تسع كتائب، وست وعشرين سرية موزعة على عمق "
            "من الساحل إلى المؤخرة. يتركز الاختبار التشغيلي على قدرة القيادة والاحتياط على امتصاص الصدمة "
            "ثم الهجوم المضاد قبل وصول الفرقة المدرعة الأولى إلى الهدف.")

    heading(doc, "4. الهدف ناصر", 1)
    ar_para(doc,
            "حُدّد OBJ NASSER-95 كقطاع تحكم على خط أنابيب ناصر-البريقة عند عمق تشغيلي يقارب 95 كم من الساحل. "
            "يمثل الهدف نقطة تأثير عملياتي واقتصادي، لكنه ليس غاية تكتيكية مستقلة؛ قيمته تأتي من كونه رمزاً "
            "لنجاح الاستثمار العميق وإمكانية قطع أو تهديد مسار الطاقة.")
    table(doc, [
        ["العامل", "التقييم"],
        ["الأهمية", "مرتفعة: خط طاقة رئيسي ومؤشر على اختراق عميق."],
        ["الوصول", "متوسط: الصحراء تسمح بالحركة، لكن المسافة واللوجستيات تضغط على القوة المدرعة."],
        ["الإصلاحية", "مرتفعة نسبياً كهدف بنية تحتية؛ التعطيل لا يعني دائماً السيطرة المستدامة."],
        ["التمييز", "متوسط: القطاع معروف على الخرائط، لكن نقطة التحكم الدقيقة تحتاج استطلاعاً وتأكيداً."],
    ], widths=[2200, 7160])

    heading(doc, "5. نقاط الإبرار الساحلي", 1)
    table(doc, [
        ["النقطة", "الإحداثيات", "الدور التشغيلي"],
        ["BLS-1", "30.2836°ش / 19.2823°ق", "إبرار تثبيت غرباً وشغل انتباه الدفاع الساحلي."],
        ["BLS-2", "30.3084°ش / 19.3648°ق", "إبرار داعم لتوسيع رأس الشاطئ وربط الجهد الغربي بالجهد الرئيسي."],
        ["BLS-3", "30.3500°ش / 19.4500°ق", "الجهد الرئيسي على شريط رملي أعرض غرب البريقة."],
        ["BLS-4", "30.4565°ش / 19.6782°ق", "إبرار التفاف خارجي على الجناح الشرقي، بقدرة محدودة وغير مضمون كمنفذ داخلي."],
    ], widths=[1500, 2500, 5360])

    heading(doc, "6. تسلسل المحاكاة خطوة بخطوة", 1)
    ar_para(doc,
            "تعرض الصفحات التالية الإطار الزمني للتشغيل. كل صورة تمثل لقطة موقف عملياتية تشمل خط المرحلة، "
            "حالة الهدف، خسائر الطرفين، وحالة نقاط الإبرار.")
    page_break(doc)

    titles_ar = [
        "الوضع الابتدائي",
        "بدء المرحلة 1 — إنزال عناصر الاستطلاع والتأمين",
        "ذروة الحرب الإلكترونية وإنذار الفريق الأزرق",
        "بدء المرحلة 2 — هبوط موجة فرقة المشاة الآلية 4",
        "توسيع رأس الشاطئ",
        "الهجوم المضاد المحلي للفريق الأزرق",
        "تثبيت رأس الجسر وإدخال فرقة المشاة الآلية 9",
        "اندفاع فرقة المشاة الآلية 9 نحو العمق",
        "بدء المرحلة 3 — إدخال الفرقة المدرعة 1",
        "هجوم الاحتياط الأزرق المضاد",
        "الهدف ناصر متنازع عليه",
        "الفصل — استيلاء الأحمر على الهدف",
    ]

    for idx, meta in enumerate(steps):
        heading(doc, f"الخطوة {idx:02d} — {titles_ar[idx]}", 2)
        table(doc, [
            ["الوقت", "المرحلة", "خط المرحلة", "حالة الهدف", "الخسائر"],
            [
                meta["time_label"],
                meta["phase"],
                f"{meta['phase_line_km']} كم",
                meta["objective_status"],
                f"أزرق {meta['losses_cumulative']['blue_destroyed']}/39 — أحمر {meta['losses_cumulative']['red_company_equivalent']} مكافئ سرية",
            ],
        ], widths=[1200, 1600, 1500, 1800, 3260])
        ar_para(doc, meta["narrative_ar"])
        add_image(doc, DIR / f"step{idx:02d}.png", f"الشكل {idx+1}: الموقف العملياتي في الخطوة {idx:02d}.")
        if idx not in (1, 3, 5, 7, 9, 11):
            page_break(doc)

    heading(doc, "7. التحليل النهائي", 1)
    table(doc, [
        ["البند", "النتيجة"],
        ["نتيجة التشغيل", "انتصار الفريق الأحمر بالسيطرة على OBJ NASSER-95 في س+144."],
        ["خسائر الفريق الأزرق", "30 وحدة من أصل 39، مع تدمير جزء كبير من سرايا المناورة الأمامية والاحتياط حول الهدف."],
        ["خسائر الفريق الأحمر", "3 مكافئ سرية، مع تدهور في اللواء المدرع 44 ومحور الجهد الرئيسي أثناء الهجوم المضاد الأزرق."],
        ["العامل الحاسم", "تأخر الاحتياط الأزرق حتى بعد استقرار رأس الشاطئ، ما سمح للفرقة المدرعة الأولى ببدء الاستثمار."],
        ["نقطة الضعف الحمراء", "اعتماد كبير على قدرة الشاطئ وممرات الخروج؛ أي تعطيل مبكر لـ BLS-3/BLS-4 كان سيغير النتيجة."],
    ], widths=[2400, 6960])
    ar_para(doc,
            "تُظهر المحاكاة أن السيطرة على الشاطئ ليست نتيجة نهائية بحد ذاتها؛ النتيجة حُسمت عندما تحولت "
            "مناطق الإبرار إلى قاعدة صيانة قتال تسمح بإدخال القوة المدرعة. الهجوم المضاد الأزرق في س+96 "
            "أبطأ الاندفاع وأوقع خسائر، لكنه لم يسبق تشكل رأس الجسر، ولذلك لم يمنع بلوغ الهدف.")

    heading(doc, "8. المراجع", 1)
    refs = [
        "nato-map-layers.geojson — بيانات الوحدات ومناطق العمليات الأصلية.",
        "enemy.docx — خطة الفريق الأحمر COA #1 ومراحل الإبرار.",
        "Current.geojson — طبقات التحليل ومواقع الإبرار والهدف التشغيلية المستخدمة في العرض.",
        "AJP-3.2 — العقيدة المشتركة للعمليات البرية، الدفاع في العمق والاحتياط.",
        "FM 3-90 / ADP 3-90 — إطار الهجوم والدفاع ونسب التخطيط.",
        "FM 55-50 و FM 71-100-2 — اعتبارات مواقع الإبرار والشاطئ والخروج.",
        "FM 34-36 App. D — إطار CARVER لتحليل الأهداف.",
    ]
    for ref in refs:
        ar_para(doc, ref, size=10, color=MUTED, after=3)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
