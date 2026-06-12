from __future__ import annotations

from pathlib import Path
import json

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Wargame2"
REPORT = OUT_DIR / "Wargame2_Report_AR.docx"

INK = RGBColor(17, 24, 39)
BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(88, 103, 124)
LIGHT = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_bidi(table) -> None:
    tbl_pr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tbl_pr.append(bidi)


def set_para_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Arial"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Arial")
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_ar_paragraph(doc, text="", style=None, size=11, bold=False, color=INK, before=0, after=6, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph(style=style)
    set_para_rtl(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    size = 16 if level == 1 else 13 if level == 2 else 12
    before = 16 if level == 1 else 12
    after = 8 if level == 1 else 6
    p = add_ar_paragraph(doc, text, size=size, bold=True, color=BLUE, before=before, after=after)
    return p


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_bidi(table)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, text in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[c_idx])
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            set_para_rtl(p)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run_font(run, size=9.5 if r_idx else 10, bold=bool(header and r_idx == 0), color=INK)
    add_ar_paragraph(doc, "", after=4)
    return table


def add_image(doc, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    cap = add_ar_paragraph(doc, caption, size=9.5, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.keep_with_next = False


def load_step(idx: int) -> dict:
    return json.loads((OUT_DIR / f"step{idx:02d}.geojson").read_text(encoding="utf-8"))


def selected_bls_rows():
    data = json.loads((OUT_DIR / "bls_selection.geojson").read_text(encoding="utf-8"))
    rows = [["النقطة", "الإحداثيات", "الدور", "تقييم الحركة / التضاريس"]]
    for feat in data["features"]:
        props = feat["properties"]
        if not props.get("selected"):
            continue
        lon, lat = feat["geometry"]["coordinates"]
        rows.append(
            [
                props["name"],
                f"{lat:.4f}°ش / {lon:.4f}°ق",
                ROLE_AR.get(props["role"], props["role"]),
                f"درجة {props['score']}؛ {MOBILITY_AR.get(props['name'], props['mobility_summary'])}",
            ]
        )
    return rows


STEP_AR = [
    "وضع ابتدائي: الدفاع الأزرق داخل منطقة العمليات، والقوة البرمائية الحمراء في التجميع البحري.",
    "بدء المرحلة الأولى: استطلاع 401، الزوارق المسيّرة المتفجرة، والحرب الإلكترونية يفتحون صورة الإبرار.",
    "تأكيد ممرات الخروج: BLS-3 هو أفضل مخرج، وBLS-2 مقيد بقرب السبخات/الأراضي الرطبة.",
    "نزول الموجة الثقيلة من فرقة المشاة الآلية الرابعة على أربع نقاط شاطئية مصححة، مع بقاء BLS-4 محدوداً كمخرج خارجي.",
    "رأس شاطئ ضحل: BLS-1 وBLS-3 يعملان، وBLS-2 محدود، وBLS-4 صالح للتثبيت لا للعبور الثقيل.",
    "هجوم مضاد أزرق مبكر يضرب رأس الشاطئ قبل اكتمال منطقة الإسناد والصيانة القتالية.",
    "دخول فرقة المشاة الآلية التاسعة يتأخر بسبب اختناق الشاطئ ومحدودية قدرة BLS-4 الخارجية.",
    "اندفاع نحو خط الصحراء المفتوحة، لكن التقدم لا يبلغ كامل خط 40-50 كم في الموعد المخطط.",
    "بدء استثمار الفرقة المدرعة الأولى عبر BLS-3، مع اعتماد خطر على ممر خروج رئيسي واحد.",
    "احتياط أزرق يضرب ممر الاستثمار قبل وصول الأحمر إلى نطاق الهدف الحاسم.",
    "ذروة عملياتية حمراء عند حافة نطاق 80-100 كم دون قدرة كافية على حشد القوة عند الهدف.",
    "الفصل النهائي: الأحمر يحتفظ برأس شاطئ ويهدد محور الأنابيب، لكن الهدف ناصر يبقى غير محتل.",
]

ROLE_AR = {
    "West fixing landing": "إبرار تثبيت غرباً",
    "Supporting landing": "إبرار داعم",
    "Main effort": "الجهد الرئيسي",
    "Eastern fixing / envelopment": "تثبيت / التفاف شرقي",
}

MOBILITY_AR = {
    "BLS-1": "حركية محدودة لكنها قابلة للاستخدام؛ تدعم الخداع وتشتيت الدفاع.",
    "BLS-2": "حركية متوسطة؛ قرب السبخات/الأراضي الرطبة يقيّد حركة المركبات.",
    "BLS-3": "مخرج شاطئي جيد؛ قرب المدافعين يجعل القمع والسرعة عاملين حاسمين.",
    "BLS-4": "احتكاك عمراني/صناعي ومائي؛ صالح للتثبيت لا لعبور ثقيل.",
}

STATUS_AR = {
    "DORMANT": "غير مفعّل",
    "THREATENED": "مهدد",
    "CONTESTED": "متنازع عليه",
    "CAPTURED": "محتل",
    "DENIED": "ممنوع",
}

PHASE_AR = {
    "PRE-H": "قبل ساعة الصفر",
    "PHASE 1": "المرحلة الأولى",
    "PHASE 2A": "المرحلة الثانية-أ",
    "PHASE 2B": "المرحلة الثانية-ب",
    "PHASE 3": "المرحلة الثالثة",
    "RESOLUTION": "الفصل النهائي",
}

FORCE_AR = {
    "Not engaged": "لم يبدأ الاشتباك",
    "1:1 reconnaissance contact": "اشتباك استطلاع 1:1",
    "1:1 contested security fight": "اشتباك تأمين متنازع 1:1",
    "2.5:1 at selected beach sectors": "2.5:1 في قطاعات الشاطئ المختارة",
    "2:1, below decisive in the east": "2:1؛ غير حاسم في الشرق",
    "1.4:1 Blue local counterstroke": "1.4:1 لصالح هجوم أزرق محلي",
    "2:1 Red but logistics-limited": "2:1 للأحمر مع قيد لوجستي",
    "2.4:1 after local consolidation": "2.4:1 بعد تثبيت موضعي",
    "2.2:1 operational, mobility penalty": "2.2:1 عملياتياً مع عقوبة حركية",
    "1.5:1 contested corridor": "1.5:1 في ممر متنازع عليه",
    "Below 3:1 at OBJ approach": "أقل من 3:1 عند مقترب الهدف",
    "Below decisive at objective": "أقل من الحسم عند الهدف",
}


def build_report() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    header = section.header.paragraphs[0]
    set_para_rtl(header, WD_ALIGN_PARAGRAPH.RIGHT)
    r = header.add_run("ملخص عملية الإبرار البرمائي - لعبة الحرب 2")
    set_run_font(r, size=9, color=MUTED)

    add_ar_paragraph(doc, "ملخص عملية الإبرار البرمائي", size=24, bold=True, color=INK, after=4)
    add_ar_paragraph(doc, "ساحل البريقة / إجدابيا - لعبة الحرب 2", size=15, bold=True, color=MUTED, after=14)
    add_table(
        doc,
        [
            ["البند", "القيمة"],
            ["نوع التشغيل", "محاكاة عملياتية حتمية مبنية على بيانات GIS، الحركية، وتوقيت الاحتياط"],
            ["زمن التشغيل", "من ساعة الصفر إلى س+144"],
            ["الهدف النهائي", "OBJ NASSER-95 على محور خط أنابيب ناصر-البريقة"],
            ["النتيجة", "الأزرق يمنع احتلال الهدف؛ الأحمر يحتفظ برأس شاطئ ويصل إلى عمق 84 كم"],
            ["الخسائر النهائية", "الأزرق 21/39 وحدة مدمرة؛ الأحمر 6 مكافئ سرية"],
        ],
        widths=[1.55, 4.85],
    )

    add_heading(doc, "1. الغرض ونطاق التشغيل")
    add_ar_paragraph(
        doc,
        "يعرض هذا التقرير خلاصة تشغيل لعبة حرب لعملية إبرار برمائي على ساحل البريقة / إجدابيا. اعتمد التشغيل على طبقات الوحدات الأصلية، خطة العدو ذات المراحل الثلاث، وطبقات GIS المحلية لتقدير صلاحية الشواطئ والحركة بعد الإبرار.",
    )
    add_ar_paragraph(
        doc,
        "الغرض هو اختبار قدرة الفريق الأحمر على تحويل الإبرار الساحلي إلى اختراق بري عميق باتجاه OBJ NASSER-95، مقابل قدرة الدفاع الأزرق والاحتياط على تعطيل الاستثمار قبل حسم الهدف.",
    )

    add_heading(doc, "2. منهجية اختيار مواقع الإبرار")
    add_ar_paragraph(
        doc,
        "تمت إعادة اختيار أربع نقاط إبرار من خط الساحل الأصلي داخل الرسم العملياتي، ثم تقييمها بعوامل: المسافة عن المدافعين، الابتعاد عن المياه الداخلية والسبخات، الابتعاد عن العمران/الصناعة، قابلية الخروج من الشاطئ، وتوزيع الجهد على الواجهة.",
    )
    add_table(doc, selected_bls_rows(), widths=[0.9, 1.4, 1.45, 2.65])
    add_ar_paragraph(
        doc,
        "ملاحظة مهمة: لا تتوفر في ملفات GIS المحلية طبقات للطرق، الارتفاع، الميل، نوع التربة التفصيلي، الأمواج، أو الأعماق البحرية؛ لذلك تعالج المحاكاة هذه العناصر كافتراضات حركية محافظة لا كقياسات مباشرة.",
        size=10,
        color=MUTED,
    )

    add_heading(doc, "3. نظام المعركة والعوامل الحاسمة")
    add_table(
        doc,
        [
            ["العنصر", "الاستخدام في التشغيل"],
            ["استطلاع 401", "تأمين ممرات الخروج وتأكيد صلاحية BLS-3 كجهد رئيسي"],
            ["24 زورقاً مسيّراً متفجراً", "إرباك الساحل وتخفيف ضغط النيران في الساعات الأولى"],
            ["فرقة المشاة الآلية 4", "موجة الاقتحام وتأسيس رأس الشاطئ"],
            ["فرقة المشاة الآلية 9", "قوة المتابعة، لكنها تأثرت باختناق BLS-2 ومحدودية BLS-4"],
            ["الفرقة المدرعة 1", "قوة الاستثمار نحو OBJ NASSER-95، مقيدة بممر خروج موثوق واحد"],
            ["لواء المدفعية 45 / كتيبة الحرب الإلكترونية 405", "إسناد ناري وإعاقة قيادة وسيطرة في المرحلتين الأولى والثانية"],
            ["كتيبة الدفاع الكيميائي 406", "قدرة حماية واستمرار عمليات؛ لا تفترض المحاكاة استخدام NBC هجومي"],
        ],
        widths=[1.75, 4.65],
    )

    add_heading(doc, "4. تسلسل التشغيل")
    for idx in range(12):
        step = load_step(idx)
        meta = step["metadata"]
        add_heading(doc, f"الخطوة {idx:02d} - {meta['time_label']} - {PHASE_AR.get(meta['phase'], meta['phase'])}", level=2)
        add_table(
            doc,
            [
                ["خط المرحلة", "حالة الهدف", "الخسائر", "قرار التحكيم"],
                [
                    f"{meta['phase_line_km']} كم",
                    STATUS_AR.get(meta["objective_status"], meta["objective_status"]),
                    f"أزرق {meta['losses_cumulative']['blue_destroyed']}/39؛ أحمر {meta['losses_cumulative']['red_company_equivalent']} مكافئ سرية",
                    FORCE_AR.get(meta["force_ratio"], meta["force_ratio"]),
                ],
            ],
            widths=[1.2, 1.2, 2.0, 2.0],
        )
        add_ar_paragraph(doc, STEP_AR[idx])
        add_image(doc, OUT_DIR / f"step{idx:02d}.jpeg", f"الشكل {idx + 1}: الموقف العملياتي في الخطوة {idx:02d}.")

    add_heading(doc, "5. النتيجة العملياتية")
    add_ar_paragraph(
        doc,
        "تنتهي اللعبة عند س+144 بمنع الفريق الأزرق احتلال OBJ NASSER-95. حقق الأحمر رأس شاطئ فعلياً وهدد محور خط الأنابيب، لكنه لم يحول ذلك إلى سيطرة على الهدف بسبب محدودية BLS-2 وBLS-4 واعتماد الاستثمار المدرع على ممر خروج رئيسي واحد من BLS-3.",
    )
    add_table(
        doc,
        [
            ["العامل", "الحكم"],
            ["نتيجة التشغيل", "منع احتلال الهدف؛ الأحمر يصل إلى 84 كم ويبلغ الذروة قبل OBJ NASSER-95"],
            ["العامل الحاسم", "توقيت الاحتياط الأزرق ضد ممر الاستثمار قبل وصول الأحمر إلى نطاق الهدف"],
            ["نقطة القوة الحمراء", "اختيار BLS-3 كجهد رئيسي وفرز الشواطئ غير الصالحة للعبور الثقيل"],
            ["نقطة الضعف الحمراء", "اختناق الشاطئ والاعتماد على ممر واحد، مع محدودية BLS-4 كمنفذ ثقيل"],
            ["الخطر المتبقي على الأزرق", "رأس الشاطئ الأحمر مستمر ويمكن أن يهدد محور الأنابيب إذا طال زمن العملية"],
        ],
        widths=[1.7, 4.7],
    )

    add_heading(doc, "6. مصادر التشغيل")
    add_ar_paragraph(doc, "nato-map-layers.geojson - بيانات الوحدات ومناطق العمليات الأصلية.")
    add_ar_paragraph(doc, "enemy.docx - خطة الفريق الأحمر ومراحل الإبرار.")
    add_ar_paragraph(doc, "Wargame1/gis/gis/landuse.geojson - تصنيف استخدامات الأرض.")
    add_ar_paragraph(doc, "Wargame1/gis/gis/inland_water.geojson - المياه الداخلية والسبخات/الأراضي الرطبة.")
    add_ar_paragraph(doc, "Wargame1/gis/gis/water_way.geojson - المجاري المائية، ولم يظهر منها عنصر مؤثر داخل مربع التشغيل.")

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build_report()
